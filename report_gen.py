"""
report_gen.py — Professional PDF & Word reports.
Modern Accounting Firm style: indigo/purple brand, coloured section bars,
cover page, executive summary table, financial highlights, colour-coded checks.
"""
from io import BytesIO
from datetime import datetime

# ── Brand colours (RGB tuples) ─────────────────────────────────────────────────
INDIGO      = (0x3B, 0x37, 0xCC)
INDIGO_DARK = (0x2D, 0x27, 0xAA)
PURPLE      = (0x7C, 0x3A, 0xED)
INDIGO_LITE = (0xE0, 0xE7, 0xFF)
WHITE       = (255, 255, 255)
BLACK       = (0x0F, 0x17, 0x2A)
GREY        = (0x64, 0x74, 0x8B)
GREY_LITE   = (0xF1, 0xF5, 0xF9)

BADGE = {
    "critical": ((0xDC, 0x26, 0x26), (0xFF, 0xF5, 0xF5)),
    "warning":  ((0xD9, 0x77, 0x06), (0xFF, 0xFB, 0xEB)),
    "info":     ((0x4F, 0x46, 0xE5), (0xEF, 0xF6, 0xFF)),
    "pass":     ((0x16, 0xA3, 0x4A), (0xF0, 0xFD, 0xF4)),
}
BADGE_LABELS = {"critical": "CRITICAL", "warning": "WARNING",
                "info": "INFO", "pass": "PASS"}


def _fmt(v):
    try:
        return f"${float(v):,.0f}" if v else "$0"
    except Exception:
        return str(v)


def _chg(cy, py):
    try:
        cy, py = float(cy or 0), float(py or 0)
        if not py:
            return "N/A"
        pct = (cy - py) / abs(py) * 100
        return f"{pct:+.1f}%"
    except Exception:
        return "N/A"


def _safe(text):
    """Strip chars that break PDF latin-1 encoding."""
    if not text:
        return ""
    text = str(text)
    replacements = {
        "\u2014": "--", "\u2013": "-",  "\u2018": "'",  "\u2019": "'",
        "\u201c": '"',  "\u201d": '"',  "\u2022": "-",  "\u2026": "...",
        "\u25b8": ">",  "\u2713": "v",  "\u00a0": " ",  "\u2192": "->",
        "\u00ae": "(R)", "\u00a9": "(c)",
    }
    for ch, rep in replacements.items():
        text = text.replace(ch, rep)
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ══════════════════════════════════════════════════════════════════════════════
#  PDF  (ReportLab)
# ══════════════════════════════════════════════════════════════════════════════
def generate_pdf(data, checks, ai_results):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import Color
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, KeepTogether, Flowable
    )
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate
    from reportlab.platypus.doctemplate import NextPageTemplate

    W, H = letter

    def rgb(t):
        return Color(t[0] / 255, t[1] / 255, t[2] / 255)

    C_IND   = rgb(INDIGO)
    C_IND_D = rgb(INDIGO_DARK)
    C_PURP  = rgb(PURPLE)
    C_IND_L = rgb(INDIGO_LITE)
    C_BLK   = rgb(BLACK)
    C_GRY   = rgb(GREY)
    C_GLITE = rgb(GREY_LITE)
    C_WHITE = Color(1, 1, 1)

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    tbl_hdr = S("th", fontName="Helvetica-Bold", fontSize=9,
                textColor=C_WHITE, leading=12, alignment=TA_CENTER)
    tbl_cell = S("tc", fontName="Helvetica", fontSize=9,
                 textColor=C_BLK, leading=12)
    tbl_cell_r = S("tcr", fontName="Helvetica", fontSize=9,
                   textColor=C_BLK, leading=12, alignment=TA_RIGHT)
    tbl_lbl = S("tl", fontName="Helvetica-Bold", fontSize=9,
                textColor=C_BLK, leading=12)
    ai_h1 = S("ah1", fontName="Helvetica-Bold", fontSize=10,
               textColor=C_IND, leading=14, spaceBefore=6, spaceAfter=3)
    ai_h2 = S("ah2", fontName="Helvetica-Bold", fontSize=9,
               textColor=C_BLK, leading=13, spaceBefore=4, spaceAfter=2)
    ai_body = S("ab", fontName="Helvetica", fontSize=9,
                textColor=C_BLK, leading=14, spaceAfter=2)
    bullet_s = S("bul", fontName="Helvetica", fontSize=9,
                 textColor=C_BLK, leading=13, leftIndent=14, spaceAfter=2)
    body_sm = S("bsm", fontName="Helvetica", fontSize=8,
                textColor=C_GRY, leading=12, leftIndent=14, spaceAfter=1)

    # client metadata
    client  = _safe(data.get("client_name", ""))
    yr_end  = str(data.get("year_end", ""))[:10]
    prep_by = _safe(data.get("prepared_by", ""))
    signer  = _safe(data.get("signer", ""))
    gen_dt  = datetime.now().strftime("%B %d, %Y")
    cy      = int(data.get("cy", 0))
    py      = int(data.get("py", 0))

    buf = BytesIO()

    def on_page(canvas, doc):
        if doc.page == 1:
            return
        canvas.saveState()
        canvas.setFillColor(C_IND_D)
        canvas.rect(0, H - 0.42 * inch, W, 0.42 * inch, fill=1, stroke=0)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(C_WHITE)
        canvas.drawString(0.5 * inch, H - 0.27 * inch,
                          f"BOOKKEEPING REVIEW  |  {client}")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(W - 0.5 * inch, H - 0.27 * inch,
                               f"Year ended {yr_end}")
        canvas.setFillColor(C_GLITE)
        canvas.rect(0, 0, W, 0.33 * inch, fill=1, stroke=0)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(C_GRY)
        canvas.drawString(0.5 * inch, 0.11 * inch,
                          "CONFIDENTIAL -- For internal review purposes only")
        canvas.drawRightString(W - 0.5 * inch, 0.11 * inch,
                               f"Page {doc.page}  |  Generated {gen_dt}")
        canvas.restoreState()

    doc_obj = BaseDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.55 * inch, bottomMargin=0.48 * inch,
    )
    frame_cover = Frame(0, 0, W, H, leftPadding=0, rightPadding=0,
                        topPadding=0, bottomPadding=0)
    frame_body = Frame(
        0.6 * inch, 0.48 * inch,
        W - 1.2 * inch, H - 1.05 * inch,
        leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
    )
    doc_obj.addPageTemplates([
        PageTemplate(id="Cover", frames=[frame_cover]),
        PageTemplate(id="Body",  frames=[frame_body], onPage=on_page),
    ])

    story = []
    BW = W - 1.2 * inch  # body width

    # ── Cover ──────────────────────────────────────────────────────────────
    class CoverPage(Flowable):
        def draw(self):
            c = self.canv
            c.setFillColor(C_IND_D)
            c.rect(0, 0, W, H, fill=1, stroke=0)
            c.setFillColor(C_PURP)
            c.rect(0, H - 0.2 * inch, W, 0.2 * inch, fill=1, stroke=0)
            # Card
            pad = 0.8 * inch
            card_y, card_h = 2.0 * inch, H - 4.3 * inch
            c.setFillColor(Color(1, 1, 1, 0.06))
            c.roundRect(pad, card_y, W - 2 * pad, card_h, 14, fill=1, stroke=0)
            cx = W / 2
            # Circle icon
            cy2 = card_y + card_h - 0.85 * inch
            c.setFillColor(C_PURP)
            c.circle(cx, cy2, 0.38 * inch, fill=1, stroke=0)
            c.setFillColor(C_WHITE)
            c.setFont("Helvetica-Bold", 18)
            c.drawCentredString(cx, cy2 - 0.07 * inch, "R")
            # Titles
            c.setFont("Helvetica-Bold", 24)
            c.setFillColor(C_WHITE)
            c.drawCentredString(cx, card_y + card_h - 1.75 * inch,
                                "BOOKKEEPING REVIEW")
            c.setFont("Helvetica", 13)
            c.setFillColor(Color(1, 1, 1, 0.7))
            c.drawCentredString(cx, card_y + card_h - 2.15 * inch, "REPORT")
            # Divider
            c.setStrokeColor(Color(1, 1, 1, 0.2))
            c.setLineWidth(0.5)
            c.line(cx - 1.4 * inch, card_y + card_h - 2.45 * inch,
                   cx + 1.4 * inch, card_y + card_h - 2.45 * inch)
            # Client
            c.setFont("Helvetica-Bold", 15)
            c.setFillColor(C_WHITE)
            c.drawCentredString(cx, card_y + card_h - 2.9 * inch, client)
            c.setFont("Helvetica", 10)
            c.setFillColor(Color(1, 1, 1, 0.7))
            c.drawCentredString(cx, card_y + card_h - 3.22 * inch,
                                f"Year Ended: {yr_end}")
            # Meta
            meta_y = card_y + 0.5 * inch
            items  = [("Prepared by", prep_by or "--"),
                      ("Reviewer", signer or "--"),
                      ("Date", gen_dt)]
            col_w  = (W - 2 * pad) / len(items)
            for i, (lbl, val) in enumerate(items):
                x = pad + i * col_w + col_w / 2
                c.setFont("Helvetica", 7.5)
                c.setFillColor(Color(1, 1, 1, 0.45))
                c.drawCentredString(x, meta_y + 0.2 * inch, lbl.upper())
                c.setFont("Helvetica-Bold", 10)
                c.setFillColor(C_WHITE)
                c.drawCentredString(x, meta_y, val)
            # Bottom
            c.setFillColor(Color(0, 0, 0, 0.3))
            c.rect(0, 0, W, 0.5 * inch, fill=1, stroke=0)
            c.setFont("Helvetica", 7.5)
            c.setFillColor(Color(1, 1, 1, 0.45))
            c.drawCentredString(W / 2, 0.18 * inch,
                "CONFIDENTIAL -- Prepared for internal CPA review purposes only")

        def wrap(self, *args):
            return (W, H)

    story.append(CoverPage())
    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())

    # ── Section bar helper ─────────────────────────────────────────────────
    def section_bar(title, fill=INDIGO_DARK):
        class Bar(Flowable):
            def __init__(self, t, f):
                self.t, self.f = t, f
            def draw(self):
                self.canv.setFillColor(rgb(self.f))
                self.canv.roundRect(0, 0, BW, 0.28 * inch, 4, fill=1, stroke=0)
                self.canv.setFillColor(C_WHITE)
                self.canv.setFont("Helvetica-Bold", 10)
                self.canv.drawString(10, 0.075 * inch, self.t.upper())
            def wrap(self, *args):
                return (BW, 0.30 * inch)
        return Bar(title, fill)

    # ── EXECUTIVE SUMMARY ──────────────────────────────────────────────────
    from auto_checks import summarize_checks
    counts = summarize_checks(checks)

    story.append(section_bar("Executive Summary"))
    story.append(Spacer(1, 0.1 * inch))

    metrics = [
        ("Total Revenue",     data.get("total_revenue_cy", 0),     data.get("total_revenue_py", 0)),
        ("Total Expenses",    data.get("total_expenses_cy", 0),    data.get("total_expenses_py", 0)),
        ("Net Income",        data.get("net_income_cy", 0),        data.get("net_income_py", 0)),
        ("Income Before Tax", data.get("income_before_tax_cy", 0), data.get("income_before_tax_py", 0)),
        ("Total Assets",      data.get("total_assets_cy", 0),      data.get("total_assets_py", 0)),
        ("Shareholder Loan",  data.get("sh_loan_cy", 0),           data.get("sh_loan_py", 0)),
        ("Retained Earnings", data.get("retained_earnings_cy", 0), data.get("retained_earnings_py", 0)),
        ("Tax Provision",     data.get("tax_provision_cy", 0),     data.get("tax_provision_py", 0)),
    ]

    hdr_row = [Paragraph(t, tbl_hdr) for t in
               ["Metric", f"CY {cy}", f"PY {py}", "Change"]]
    rows = [hdr_row]
    for i, (lbl, cv, pv) in enumerate(metrics):
        ch = _chg(cv, pv)
        chg_c = rgb((0x16, 0xA3, 0x4A)) if ch.startswith("+") else (
                rgb((0xDC, 0x26, 0x26)) if ch.startswith("-") else C_GRY)
        rows.append([
            Paragraph(lbl, tbl_lbl),
            Paragraph(_fmt(cv), tbl_cell_r),
            Paragraph(_fmt(pv), tbl_cell_r),
            Paragraph(ch, S("chg", fontName="Helvetica-Bold", fontSize=9,
                             textColor=chg_c, leading=12, alignment=TA_RIGHT)),
        ])

    mt = Table(rows, colWidths=[BW * 0.38, BW * 0.22, BW * 0.22, BW * 0.18])
    mt.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), C_IND),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [C_GLITE, C_WHITE]),
        ("GRID",          (0, 0), (-1, -1), 0.3, rgb((0xE2, 0xE8, 0xF0))),
        ("ALIGN",         (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN",         (0, 0), (0, -1),  "LEFT"),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(mt)
    story.append(Spacer(1, 0.1 * inch))

    # Status bar
    crit = counts["critical"]
    warn = counts["warning"]
    if crit == 0 and warn == 0:
        st_txt, st_c = "READY TO CLIENT", (0x16, 0xA3, 0x4A)
    elif crit == 0:
        st_txt, st_c = "MINOR CORRECTIONS NEEDED", (0xD9, 0x77, 0x06)
    else:
        st_txt, st_c = "MAJOR CORRECTIONS NEEDED", (0xDC, 0x26, 0x26)

    class StatusBar(Flowable):
        def draw(self):
            self.canv.setFillColor(rgb(st_c))
            self.canv.roundRect(0, 0, BW, 0.28 * inch, 8, fill=1, stroke=0)
            self.canv.setFillColor(C_WHITE)
            self.canv.setFont("Helvetica-Bold", 10)
            self.canv.drawCentredString(BW / 2, 0.075 * inch,
                                        f"REVIEW STATUS:  {st_txt}")
        def wrap(self, *args):
            return (BW, 0.30 * inch)

    story.append(StatusBar())
    story.append(Spacer(1, 0.1 * inch))

    # Check count grid
    badge_data = [[
        Paragraph(f"<b>{counts['critical']}</b><br/>Critical",
                  S("", fontName="Helvetica-Bold", fontSize=15,
                    textColor=rgb(BADGE["critical"][0]), leading=18, alignment=TA_CENTER)),
        Paragraph(f"<b>{counts['warning']}</b><br/>Warnings",
                  S("", fontName="Helvetica-Bold", fontSize=15,
                    textColor=rgb(BADGE["warning"][0]), leading=18, alignment=TA_CENTER)),
        Paragraph(f"<b>{counts['info']}</b><br/>Info",
                  S("", fontName="Helvetica-Bold", fontSize=15,
                    textColor=rgb(BADGE["info"][0]), leading=18, alignment=TA_CENTER)),
        Paragraph(f"<b>{counts['pass']}</b><br/>Passed",
                  S("", fontName="Helvetica-Bold", fontSize=15,
                    textColor=rgb(BADGE["pass"][0]), leading=18, alignment=TA_CENTER)),
    ]]
    cw2 = BW / 4
    ct = Table(badge_data, colWidths=[cw2] * 4)
    ct.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (0, 0), rgb(BADGE["critical"][1])),
        ("BACKGROUND",    (1, 0), (1, 0), rgb(BADGE["warning"][1])),
        ("BACKGROUND",    (2, 0), (2, 0), rgb(BADGE["info"][1])),
        ("BACKGROUND",    (3, 0), (3, 0), rgb(BADGE["pass"][1])),
        ("BOX",           (0, 0), (-1, -1), 0.5, rgb((0xE2, 0xE8, 0xF0))),
        ("INNERGRID",     (0, 0), (-1, -1), 0.5, rgb((0xE2, 0xE8, 0xF0))),
        ("TOPPADDING",    (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(ct)
    story.append(Spacer(1, 0.2 * inch))

    # ── FINANCIAL HIGHLIGHTS ───────────────────────────────────────────────
    story.append(section_bar("Financial Highlights", PURPLE))
    story.append(Spacer(1, 0.1 * inch))

    def fin_table_pdf(label, items, total_cy, total_py, hdr_colour):
        if not items:
            return []
        hdr = [Paragraph(t, tbl_hdr) for t in
               [label, f"CY {cy}", f"PY {py}", "Change"]]
        rows = [hdr]
        for i, (lbl2, vals) in enumerate(list(items.items())[:20]):
            cv2 = vals.get("cy", 0)
            pv2 = vals.get("py", 0)
            if not cv2 and not pv2:
                continue
            ch = _chg(cv2, pv2)
            chg_c = rgb((0x16, 0xA3, 0x4A)) if ch.startswith("+") else (
                    rgb((0xDC, 0x26, 0x26)) if ch.startswith("-") else C_GRY)
            rows.append([
                Paragraph(_safe(lbl2), tbl_cell),
                Paragraph(_fmt(cv2), tbl_cell_r),
                Paragraph(_fmt(pv2), tbl_cell_r),
                Paragraph(ch, S("", fontName="Helvetica-Bold", fontSize=9,
                                 textColor=chg_c, leading=12, alignment=TA_RIGHT)),
            ])
        # Total
        ch_tot = _chg(total_cy, total_py)
        rows.append([
            Paragraph("TOTAL", S("", fontName="Helvetica-Bold", fontSize=9,
                                  textColor=C_WHITE, leading=12)),
            Paragraph(_fmt(total_cy), S("", fontName="Helvetica-Bold", fontSize=9,
                                         textColor=C_WHITE, leading=12, alignment=TA_RIGHT)),
            Paragraph(_fmt(total_py), S("", fontName="Helvetica-Bold", fontSize=9,
                                         textColor=C_WHITE, leading=12, alignment=TA_RIGHT)),
            Paragraph(ch_tot, S("", fontName="Helvetica-Bold", fontSize=9,
                                  textColor=C_WHITE, leading=12, alignment=TA_RIGHT)),
        ])
        t = Table(rows, colWidths=[BW * 0.40, BW * 0.22, BW * 0.20, BW * 0.18])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0),  (-1, 0),  rgb(hdr_colour)),
            ("BACKGROUND",    (0, -1), (-1, -1), C_IND),
            ("ROWBACKGROUNDS",(0, 1),  (-1, -2), [C_GLITE, C_WHITE]),
            ("GRID",          (0, 0),  (-1, -1), 0.25, rgb((0xE2, 0xE8, 0xF0))),
            ("ALIGN",         (1, 0),  (-1, -1), "RIGHT"),
            ("ALIGN",         (0, 0),  (0, -1),  "LEFT"),
            ("LEFTPADDING",   (0, 0),  (-1, -1), 8),
            ("RIGHTPADDING",  (0, 0),  (-1, -1), 8),
            ("TOPPADDING",    (0, 0),  (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 4),
        ]))
        return [t, Spacer(1, 0.1 * inch)]

    story += fin_table_pdf("Revenue", data.get("income_items", {}),
                           data.get("total_revenue_cy", 0), data.get("total_revenue_py", 0),
                           PURPLE)
    story += fin_table_pdf("Expenses", data.get("expense_items", {}),
                           data.get("total_expenses_cy", 0), data.get("total_expenses_py", 0),
                           INDIGO)
    story.append(Spacer(1, 0.1 * inch))

    # ── AUTOMATIC CHECKS ───────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(section_bar("Automatic Checks"))
    story.append(Spacer(1, 0.1 * inch))

    cats = {}
    for c in checks:
        cats.setdefault(c.get("category", "General"), []).append(c)
    level_order = {"critical": 0, "warning": 1, "info": 2, "pass": 3}

    for cat, cat_checks in cats.items():
        class CatBar(Flowable):
            def __init__(self, t):
                self.t = t
            def draw(self):
                self.canv.setFillColor(C_IND_L)
                self.canv.rect(0, 0, BW, 0.21 * inch, fill=1, stroke=0)
                self.canv.setFillColor(C_IND)
                self.canv.setFont("Helvetica-Bold", 8)
                self.canv.drawString(8, 0.055 * inch, f"  {self.t.upper()}")
            def wrap(self, *args):
                return (BW, 0.23 * inch)

        story.append(KeepTogether([CatBar(cat), Spacer(1, 0.04 * inch)]))

        for chk in sorted(cat_checks, key=lambda x: level_order.get(x["level"], 9)):
            lv = chk["level"]
            bdg_c  = rgb(BADGE[lv][0])
            bg_c   = rgb(BADGE[lv][1])
            lbl    = BADGE_LABELS.get(lv, lv.upper())
            t_text = _safe(chk.get("title", ""))
            d_text = _safe(chk.get("detail", ""))

            badge_p = Paragraph(lbl, S("",
                fontName="Helvetica-Bold", fontSize=7.5,
                textColor=C_WHITE, leading=10, alignment=TA_CENTER))
            title_p = Paragraph(t_text, S("",
                fontName="Helvetica-Bold", fontSize=9,
                textColor=C_BLK, leading=12))

            row_data = [[badge_p, title_p]]
            rt = Table(row_data, colWidths=[BW * 0.13, BW * 0.87])
            rt.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, 0), bdg_c),
                ("BACKGROUND",    (1, 0), (1, 0), bg_c),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING",    (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING",   (0, 0), (-1, -1), 8),
                ("BOX",           (0, 0), (-1, -1), 0.3, rgb((0xE2, 0xE8, 0xF0))),
            ]))
            keep = [rt]
            if d_text and lv != "pass":
                det_p = Paragraph(d_text, body_sm)
                dt = Table([[det_p]], colWidths=[BW])
                dt.setStyle(TableStyle([
                    ("BACKGROUND",    (0, 0), (-1, -1), bg_c),
                    ("LEFTPADDING",   (0, 0), (-1, -1), 22),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING",    (0, 0), (-1, -1), 1),
                    ("BOX",           (0, 0), (-1, -1), 0.3, rgb((0xE2, 0xE8, 0xF0))),
                ]))
                keep.append(dt)
            keep.append(Spacer(1, 0.03 * inch))
            story.append(KeepTogether(keep))

        story.append(Spacer(1, 0.06 * inch))

    # ── AI SECTIONS ────────────────────────────────────────────────────────
    PROMPT_TITLES = {
        "full_review":        "Full File Review",
        "tax_planning":       "Tax Planning Opportunities",
        "missing_expenses":   "Missing Expenses Analysis",
        "staff_queries":      "Staff Queries & Corrections",
        "management_summary": "Management Summary (Internal)",
        "client_summary":     "Client Executive Summary",
        "engagement_notes":   "Engagement File Notes",
        "unusual_items":         "Unusual Items -- CRA Risk Flags",
        "bank_statement_review": "Bank Statement Audit Review",
    }
    SEC_FILLS = [INDIGO_DARK, PURPLE, INDIGO_DARK, PURPLE,
                 INDIGO_DARK, PURPLE, INDIGO_DARK, PURPLE, INDIGO_DARK]

    for si, (key, content) in enumerate(ai_results.items()):
        if not content:
            continue
        story.append(PageBreak())
        title_str = PROMPT_TITLES.get(key, key.replace("_", " ").title())
        story.append(section_bar(title_str, SEC_FILLS[si % len(SEC_FILLS)]))
        story.append(Spacer(1, 0.12 * inch))

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.05 * inch))
                continue
            line = _safe(line)
            if line.startswith("### "):
                story.append(Paragraph(line[4:], ai_h2))
            elif line.startswith("## ") or line.startswith("# "):
                story.append(Paragraph(line.lstrip("#").strip(), ai_h1))
            elif line.startswith("**") and line.endswith("**") and len(line) > 4:
                story.append(Paragraph(line.replace("**", ""), ai_h2))
            elif line.isupper() and 4 < len(line) < 70:
                story.append(Paragraph(line, ai_h1))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"&bull;  {line[2:]}", bullet_s))
            elif len(line) > 2 and line[0].isdigit() and line[1] in ".):":
                story.append(Paragraph(line, S("",
                    fontName="Helvetica", fontSize=9, textColor=C_BLK,
                    leading=14, leftIndent=14, spaceAfter=2)))
            else:
                story.append(Paragraph(line, ai_body))

    doc_obj.build(story)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
#  WORD  (python-docx)
# ══════════════════════════════════════════════════════════════════════════════
def generate_word(data, checks, ai_results):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from auto_checks import summarize_checks

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin   = Inches(0.85)
        sec.right_margin  = Inches(0.85)

    def rc(t):
        return RGBColor(t[0], t[1], t[2])

    def shade(cell, colour_tuple):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        hex_col = "{:02X}{:02X}{:02X}".format(*colour_tuple)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_col)
        tcPr.append(shd)

    def cp(cell, text, bold=False, size=9, color=None,
           align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(str(text))
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    def section_heading(title, colour=INDIGO_DARK):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(5)
        run = p.add_run(f"  {title.upper()}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*colour))
        pPr.append(shd)

    # ── COVER ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("BOOKKEEPING REVIEW REPORT")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = rc(INDIGO_DARK)

    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_after = Pt(20)
    dr = div.add_run("_" * 44)
    dr.font.color.rgb = rc(PURPLE)
    dr.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(_safe(data.get("client_name", "")))
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = rc(BLACK)

    yr_p = doc.add_paragraph()
    yr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yr_p.paragraph_format.space_after = Pt(30)
    yr_r = yr_p.add_run(f"Year Ended: {str(data.get('year_end',''))[:10]}")
    yr_r.font.size = Pt(12)
    yr_r.font.color.rgb = rc(GREY)

    # Meta table
    mt = doc.add_table(rows=2, cols=3)
    mt.style = "Table Grid"
    for i, (h, v) in enumerate([
        ("Prepared by", _safe(data.get("prepared_by", "--"))),
        ("Reviewer / Signer", _safe(data.get("signer", "--"))),
        ("Date Generated", datetime.now().strftime("%B %d, %Y")),
    ]):
        shade(mt.rows[0].cells[i], INDIGO_DARK)
        shade(mt.rows[1].cells[i], INDIGO_LITE)
        cp(mt.rows[0].cells[i], h, bold=True, size=9,
           color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
        cp(mt.rows[1].cells[i], v, bold=True, size=10,
           color=rc(INDIGO), align=WD_ALIGN_PARAGRAPH.CENTER)
        mt.rows[0].cells[i].width = Inches(2.1)
        mt.rows[1].cells[i].width = Inches(2.1)

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ──────────────────────────────────────────────────
    section_heading("Executive Summary")
    counts = summarize_checks(checks)
    cy = int(data.get("cy", 0))
    py = int(data.get("py", 0))

    metrics = [
        ("Total Revenue",     data.get("total_revenue_cy", 0),     data.get("total_revenue_py", 0)),
        ("Total Expenses",    data.get("total_expenses_cy", 0),    data.get("total_expenses_py", 0)),
        ("Net Income",        data.get("net_income_cy", 0),        data.get("net_income_py", 0)),
        ("Income Before Tax", data.get("income_before_tax_cy", 0), data.get("income_before_tax_py", 0)),
        ("Total Assets",      data.get("total_assets_cy", 0),      data.get("total_assets_py", 0)),
        ("Shareholder Loan",  data.get("sh_loan_cy", 0),           data.get("sh_loan_py", 0)),
        ("Retained Earnings", data.get("retained_earnings_cy", 0), data.get("retained_earnings_py", 0)),
        ("Tax Provision",     data.get("tax_provision_cy", 0),     data.get("tax_provision_py", 0)),
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Metric", f"CY {cy}", f"PY {py}", "Change"]):
        shade(tbl.rows[0].cells[i], INDIGO)
        cp(tbl.rows[0].cells[i], h, bold=True, size=9,
           color=RGBColor(255, 255, 255),
           align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)

    for idx, (lbl, cv, pv) in enumerate(metrics):
        bg = GREY_LITE if idx % 2 == 0 else WHITE
        ch = _chg(cv, pv)
        ch_c = rc((0x16, 0xA3, 0x4A)) if ch.startswith("+") else (
               rc((0xDC, 0x26, 0x26)) if ch.startswith("-") else rc(GREY))
        row = tbl.add_row()
        for ci, (txt, aln, clr, bld) in enumerate([
            (lbl,      WD_ALIGN_PARAGRAPH.LEFT,  rc(BLACK), True),
            (_fmt(cv), WD_ALIGN_PARAGRAPH.RIGHT, rc(BLACK), False),
            (_fmt(pv), WD_ALIGN_PARAGRAPH.RIGHT, rc(GREY),  False),
            (ch,       WD_ALIGN_PARAGRAPH.RIGHT, ch_c,      True),
        ]):
            shade(row.cells[ci], bg)
            cp(row.cells[ci], txt, bold=bld, size=9, color=clr, align=aln)

    for row in tbl.rows:
        for i, w in enumerate([Inches(2.5), Inches(1.4), Inches(1.4), Inches(1.1)]):
            row.cells[i].width = w

    doc.add_paragraph()

    # Status bar
    if counts["critical"] == 0 and counts["warning"] == 0:
        st_txt, st_c = "READY TO CLIENT", (0x16, 0xA3, 0x4A)
    elif counts["critical"] == 0:
        st_txt, st_c = "MINOR CORRECTIONS NEEDED", (0xD9, 0x77, 0x06)
    else:
        st_txt, st_c = "MAJOR CORRECTIONS NEEDED", (0xDC, 0x26, 0x26)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(10)
    sr = sp.add_run(f"REVIEW STATUS:  {st_txt}")
    sr.bold = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(255, 255, 255)
    pPr = sp._p.get_or_add_pPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear")
    shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*st_c))
    pPr.append(shd2)

    doc.add_page_break()

    # ── FINANCIAL HIGHLIGHTS ───────────────────────────────────────────────
    section_heading("Financial Highlights", PURPLE)
    doc.add_paragraph()

    def fin_word_tbl(label, items, total_cy, total_py, hdr_colour):
        if not items:
            return
        lp = doc.add_paragraph()
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = rc(hdr_colour)
        lp.paragraph_format.space_after = Pt(2)

        ft = doc.add_table(rows=1, cols=4)
        ft.style = "Table Grid"
        for i, h in enumerate([label, f"CY {cy}", f"PY {py}", "Change"]):
            shade(ft.rows[0].cells[i], hdr_colour)
            cp(ft.rows[0].cells[i], h, bold=True, size=8,
               color=RGBColor(255, 255, 255),
               align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        for idx, (lbl2, vals) in enumerate(list(items.items())[:20]):
            cv2, pv2 = vals.get("cy", 0), vals.get("py", 0)
            if not cv2 and not pv2:
                continue
            ch = _chg(cv2, pv2)
            ch_c = rc((0x16, 0xA3, 0x4A)) if ch.startswith("+") else (
                   rc((0xDC, 0x26, 0x26)) if ch.startswith("-") else rc(GREY))
            bg = GREY_LITE if idx % 2 == 0 else WHITE
            row = ft.add_row()
            for ci, (txt, aln, clr, bld) in enumerate([
                (_safe(lbl2), WD_ALIGN_PARAGRAPH.LEFT,  rc(BLACK), False),
                (_fmt(cv2),   WD_ALIGN_PARAGRAPH.RIGHT, rc(BLACK), False),
                (_fmt(pv2),   WD_ALIGN_PARAGRAPH.RIGHT, rc(GREY),  False),
                (ch,          WD_ALIGN_PARAGRAPH.RIGHT, ch_c,      True),
            ]):
                shade(row.cells[ci], bg)
                cp(row.cells[ci], txt, bold=bld, size=8, color=clr, align=aln)
        tot_row = ft.add_row()
        ch_t = _chg(total_cy, total_py)
        for ci, (txt, aln) in enumerate([
            ("TOTAL",        WD_ALIGN_PARAGRAPH.LEFT),
            (_fmt(total_cy), WD_ALIGN_PARAGRAPH.RIGHT),
            (_fmt(total_py), WD_ALIGN_PARAGRAPH.RIGHT),
            (ch_t,           WD_ALIGN_PARAGRAPH.RIGHT),
        ]):
            shade(tot_row.cells[ci], INDIGO)
            cp(tot_row.cells[ci], txt, bold=True, size=8,
               color=RGBColor(255, 255, 255), align=aln)
        for row in ft.rows:
            for i, w in enumerate([Inches(2.5), Inches(1.4), Inches(1.4), Inches(1.1)]):
                row.cells[i].width = w
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    fin_word_tbl("Revenue",  data.get("income_items", {}),
                 data.get("total_revenue_cy", 0), data.get("total_revenue_py", 0), PURPLE)
    fin_word_tbl("Expenses", data.get("expense_items", {}),
                 data.get("total_expenses_cy", 0), data.get("total_expenses_py", 0), INDIGO)

    doc.add_page_break()

    # ── AUTOMATIC CHECKS ───────────────────────────────────────────────────
    section_heading("Automatic Checks")
    doc.add_paragraph()

    cats = {}
    for c in checks:
        cats.setdefault(c.get("category", "General"), []).append(c)
    level_order = {"critical": 0, "warning": 1, "info": 2, "pass": 3}

    for cat, cat_checks in cats.items():
        cat_p = doc.add_paragraph()
        cat_p.paragraph_format.space_before = Pt(8)
        cat_p.paragraph_format.space_after  = Pt(3)
        cat_r = cat_p.add_run(f"  {cat.upper()}")
        cat_r.bold = True
        cat_r.font.size = Pt(8.5)
        cat_r.font.color.rgb = rc(INDIGO)
        pPr = cat_p._p.get_or_add_pPr()
        shd3 = OxmlElement("w:shd")
        shd3.set(qn("w:val"), "clear")
        shd3.set(qn("w:color"), "auto")
        shd3.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*INDIGO_LITE))
        pPr.append(shd3)

        for chk in sorted(cat_checks, key=lambda x: level_order.get(x["level"], 9)):
            lv = chk["level"]
            lbl = BADGE_LABELS.get(lv, lv.upper())
            ctbl = doc.add_table(rows=1, cols=2)
            ctbl.style = "Table Grid"
            shade(ctbl.rows[0].cells[0], BADGE[lv][0])
            shade(ctbl.rows[0].cells[1], BADGE[lv][1])
            ctbl.rows[0].cells[0].width = Inches(0.85)
            ctbl.rows[0].cells[1].width = Inches(5.6)
            cp(ctbl.rows[0].cells[0], lbl, bold=True, size=8,
               color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
            cp(ctbl.rows[0].cells[1], _safe(chk.get("title", "")),
               bold=True, size=9, color=rc(BLACK))
            if chk.get("detail") and lv != "pass":
                det = ctbl.add_row()
                shade(det.cells[0], BADGE[lv][0])
                shade(det.cells[1], BADGE[lv][1])
                cp(det.cells[0], "", size=8)
                cp(det.cells[1], _safe(chk.get("detail", "")),
                   size=8, color=rc(GREY))
            pp = doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(1)

    doc.add_page_break()

    # ── AI SECTIONS ────────────────────────────────────────────────────────
    PROMPT_TITLES = {
        "full_review":        "Full File Review",
        "tax_planning":       "Tax Planning Opportunities",
        "missing_expenses":   "Missing Expenses Analysis",
        "staff_queries":      "Staff Queries & Corrections",
        "management_summary": "Management Summary (Internal)",
        "client_summary":     "Client Executive Summary",
        "engagement_notes":   "Engagement File Notes",
        "unusual_items":         "Unusual Items -- CRA Risk Flags",
        "bank_statement_review": "Bank Statement Audit Review",
    }
    SEC_FILLS = [INDIGO_DARK, PURPLE, INDIGO_DARK, PURPLE,
                 INDIGO_DARK, PURPLE, INDIGO_DARK, PURPLE, INDIGO_DARK]

    for si, (key, content) in enumerate(ai_results.items()):
        if not content:
            continue
        title_str = PROMPT_TITLES.get(key, key.replace("_", " ").title())
        section_heading(title_str, SEC_FILLS[si % len(SEC_FILLS)])
        doc.add_paragraph()

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue
            line = _safe(line)
            if line.startswith("### "):
                p = doc.add_paragraph(line[4:])
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = rc(BLACK)
                p.paragraph_format.space_before = Pt(4)
            elif line.startswith("## ") or line.startswith("# "):
                p = doc.add_paragraph(line.lstrip("#").strip())
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = rc(INDIGO)
                p.paragraph_format.space_before = Pt(6)
            elif line.startswith("**") and line.endswith("**") and len(line) > 4:
                p = doc.add_paragraph(line.replace("**", ""))
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = rc(INDIGO)
            elif line.isupper() and 4 < len(line) < 70:
                p = doc.add_paragraph(line)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = rc(INDIGO)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(line[2:]).font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.2)
            elif len(line) > 2 and line[0].isdigit() and line[1] in ".):":
                p = doc.add_paragraph(style="List Number")
                p.add_run(line).font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.2)
            else:
                p = doc.add_paragraph(line)
                if p.runs:
                    p.runs[0].font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()
        if si < len(ai_results) - 1:
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
