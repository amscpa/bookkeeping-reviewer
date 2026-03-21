"""
pages/2_YoY_Comparison.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Year-over-Year Financial Statement Comparison — STANDALONE PAGE
• Uses OpenAI o3 (reasoning / "thinking" model) — no temperature needed
• Generates a formatted PDF report  +  formatted Word (.docx) report
• Zero changes to any existing app file
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

# ─── Standard library ──────────────────────────────────────────────────────
import re
import textwrap
from io import BytesIO
from datetime import datetime

# ─── Third-party ───────────────────────────────────────────────────────────
import streamlit as st
import pandas as pd
from openai import OpenAI

# ─── ReportLab (PDF) ───────────────────────────────────────────────────────
from reportlab.lib import colors as rl_colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, HRFlowable, PageBreak, KeepTogether
)

# ─── python-docx (Word) ────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="YoY Comparison | Bookkeeping Reviewer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
# BRAND COLOURS  (matches your existing app navy/teal theme)
# ══════════════════════════════════════════════════════════════════════════════
NAVY      = rl_colors.HexColor("#1B2A4A")
TEAL      = rl_colors.HexColor("#2E86AB")
GOLD      = rl_colors.HexColor("#F0A500")
LIGHT_BG  = rl_colors.HexColor("#F4F7FB")
MID_GREY  = rl_colors.HexColor("#D0D7E3")
GREEN_CL  = rl_colors.HexColor("#1A7F4B")
RED_CL    = rl_colors.HexColor("#C0392B")
GREEN_BG  = rl_colors.HexColor("#E8F8F0")
RED_BG    = rl_colors.HexColor("#FDECEA")
AMBER_BG  = rl_colors.HexColor("#FFF8E1")

# Word hex strings (no #)
W_NAVY  = "1B2A4A"
W_TEAL  = "2E86AB"
W_GOLD  = "F0A500"
W_GREEN = "1A7F4B"
W_RED   = "C0392B"
W_LGREY = "F4F7FB"
W_MGREY = "D0D7E3"

# ══════════════════════════════════════════════════════════════════════════════
# EXCEL HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def get_sheets(f) -> list[str]:
    f.seek(0)
    xl = pd.ExcelFile(BytesIO(f.read()))
    f.seek(0)
    return xl.sheet_names


def read_sheet(f, sheet: str) -> pd.DataFrame:
    f.seek(0)
    df = pd.read_excel(BytesIO(f.read()), sheet_name=sheet,
                       header=None, dtype=str)
    f.seek(0)
    return df.fillna("")


def df_to_text(df: pd.DataFrame, max_rows: int = 250) -> str:
    lines = []
    for _, row in df.head(max_rows).iterrows():
        cells = [str(c).strip() for c in row if str(c).strip()]
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)

# ══════════════════════════════════════════════════════════════════════════════
# OPENAI — o3 REASONING MODEL
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM = """You are a senior Canadian CPA (CPA designation) with 20+ years experience
reviewing Alberta small-business corporation financial statements.
You are producing a formal Year-over-Year variance analysis report.
Use clear Markdown: ## for section headers, pipe tables with header rows,
**bold** for key figures, and bullet lists where appropriate.
Be precise — reference actual dollar amounts. All amounts are CAD."""


def build_prompt(cy: str, py: str,
                 cy_is: str, py_is: str,
                 cy_bs: str, py_bs: str) -> str:
    return f"""
Compare the two fiscal years below and produce a complete professional analysis.

CURRENT YEAR:  {cy}
PRIOR YEAR:    {py}

═══ INCOME STATEMENT — {cy} ═══
{cy_is}

═══ INCOME STATEMENT — {py} ═══
{py_is}

═══ BALANCE SHEET — {cy} ═══
{cy_bs}

═══ BALANCE SHEET — {py} ═══
{py_bs}

Produce EXACTLY these eight sections with these EXACT headers:

## 1. EXECUTIVE SUMMARY
4–6 sentences: overall revenue/profit trend, biggest balance sheet movements,
top risk or opportunity, and a one-sentence overall assessment.

## 2. INCOME STATEMENT VARIANCE
Markdown table, columns:
| Account | {py} | {cy} | $ Change | % Change | Flag |
Flag values: ▲ Increase | ▼ Decrease | ⚠ Unusual | → Stable
Include every revenue and expense line you can identify. Add a TOTAL NET INCOME row.

## 3. BALANCE SHEET VARIANCE
Markdown table, columns:
| Account | {py} | {cy} | $ Change | % Change | Flag |
Include every asset, liability, and equity line. Add TOTAL ASSETS, TOTAL LIABILITIES, and EQUITY rows.

## 4. KEY DRIVERS OF CHANGE
Numbered list. Top 6–8 accounts with largest dollar or % movement.
For each: account name, dollar change, % change, and 1–2 sentence business explanation.

## 5. PROFITABILITY & RATIO ANALYSIS
Table:
| Ratio | {py} | {cy} | Change | Comment |
Include: Gross Margin %, Net Margin %, Effective Tax Rate %,
Current Ratio, Debt-to-Equity, Revenue Growth %.

## 6. RED FLAGS & CRA RISK ITEMS
Numbered list. For each: account/item, issue description,
risk level (🔴 High / 🟡 Medium / 🟢 Low), recommended action.

## 7. YEAR-TO-YEAR CONSISTENCY CHECK
List any accounts that appear in one year but not the other (new or discontinued).
Note any prior-year comparative mismatches that need investigation.

## 8. CLIENT MEETING RECOMMENDATIONS
Numbered list of 4–6 specific talking points the CPA should raise with the client,
with estimated dollar impact where applicable.
""".strip()


def call_openai(prompt: str, api_key: str, model: str) -> str:
    client = OpenAI(api_key=api_key)
    # o3 / o4-mini use max_completion_tokens; gpt-4o uses max_tokens
    if model.startswith("o"):
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM},
                {"role": "user",   "content": prompt},
            ],
            max_completion_tokens=4000,
        )
    else:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM},
                {"role": "user",   "content": prompt},
            ],
            max_tokens=4000,
            temperature=0.2,
        )
    return resp.choices[0].message.content

# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ══════════════════════════════════════════════════════════════════════════════

def parse_sections(text: str) -> list[dict]:
    """Split markdown into [{title, body}] blocks."""
    sections, buf, title = [], [], "Preamble"
    for line in text.splitlines():
        if line.startswith("## "):
            if buf:
                sections.append({"title": title, "body": "\n".join(buf).strip()})
            title, buf = line[3:].strip(), []
        else:
            buf.append(line)
    if buf:
        sections.append({"title": title, "body": "\n".join(buf).strip()})
    return sections


def parse_md_table(body: str) -> list[list[str]]:
    """Extract rows from a markdown pipe table (skip separator lines)."""
    rows = []
    for line in body.splitlines():
        stripped = line.strip()
        if stripped.startswith("|") and not re.match(r"^\|[-| :]+\|$", stripped):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            rows.append(cells)
    return rows


def strip_md(text: str) -> str:
    """Remove bold/italic markers for plain-text rendering."""
    return re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", text)

# ══════════════════════════════════════════════════════════════════════════════
# FLAG COLOURS  (for table cells)
# ══════════════════════════════════════════════════════════════════════════════

def flag_color_rl(flag: str):
    if "▲" in flag:
        return GREEN_BG
    if "▼" in flag:
        return RED_BG
    if "⚠" in flag:
        return AMBER_BG
    return None


def flag_color_word(flag: str):
    if "▲" in flag:
        return "E8F8F0"
    if "▼" in flag:
        return "FDECEA"
    if "⚠" in flag:
        return "FFF8E1"
    return None

# ══════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(sections: list[dict], cy: str, py: str,
              firm_name: str, prepared_by: str) -> BytesIO:

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
    )

    # ── Styles ──────────────────────────────────────────────────────────────
    base = getSampleStyleSheet()

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sty_cover_firm = S("CoverFirm",
        fontSize=11, textColor=rl_colors.white, alignment=TA_CENTER,
        fontName="Helvetica", spaceAfter=4)
    sty_cover_title = S("CoverTitle",
        fontSize=22, textColor=rl_colors.white, alignment=TA_CENTER,
        fontName="Helvetica-Bold", spaceAfter=6)
    sty_cover_sub = S("CoverSub",
        fontSize=13, textColor=rl_colors.white, alignment=TA_CENTER,
        fontName="Helvetica", spaceAfter=4)
    sty_cover_meta = S("CoverMeta",
        fontSize=10, textColor=GOLD, alignment=TA_CENTER,
        fontName="Helvetica", spaceAfter=3)

    sty_sec_hdr = S("SecHdr",
        fontSize=12, textColor=rl_colors.white,
        fontName="Helvetica-Bold", leftIndent=6,
        spaceAfter=6, spaceBefore=14)
    sty_body = S("Body",
        fontSize=9.5, textColor=rl_colors.black,
        fontName="Helvetica", leading=14, spaceAfter=4)
    sty_bullet = S("Bullet",
        fontSize=9.5, fontName="Helvetica", leading=13,
        leftIndent=16, spaceAfter=3, bulletIndent=6)
    sty_th = S("TH",
        fontSize=9, textColor=rl_colors.white,
        fontName="Helvetica-Bold", alignment=TA_CENTER)
    sty_td = S("TD",
        fontSize=8.5, fontName="Helvetica", alignment=TA_LEFT, leading=12)
    sty_td_r = S("TDR",
        fontSize=8.5, fontName="Helvetica", alignment=TA_RIGHT, leading=12)
    sty_footer = S("Footer",
        fontSize=8, textColor=MID_GREY, alignment=TA_CENTER,
        fontName="Helvetica-Oblique")

    story = []

    # ══ COVER PAGE ════════════════════════════════════════════════════════
    # Full-width navy banner via a single-cell table
    cover_data = [[
        Paragraph(firm_name.upper(), sty_cover_firm),
    ]]
    cover_top = Table(cover_data, colWidths=[7*inch])
    cover_top.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,-1), NAVY),
        ("TOPPADDING",  (0,0), (-1,-1), 28),
        ("BOTTOMPADDING",(0,0),(-1,-1), 10),
        ("LEFTPADDING", (0,0), (-1,-1), 12),
    ]))
    story.append(cover_top)
    story.append(Spacer(1, 0.1*inch))

    # Title block
    title_data = [[
        Paragraph("YEAR-OVER-YEAR", sty_cover_title),
    ],[
        Paragraph("Financial Statement Comparison", sty_cover_sub),
    ],[
        Paragraph(f"{py}  →  {cy}", sty_cover_sub),
    ]]
    title_tbl = Table(title_data, colWidths=[7*inch])
    title_tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,-1), TEAL),
        ("TOPPADDING",  (0,0), (-1,-1), 12),
        ("BOTTOMPADDING",(0,0),(-1,-1), 12),
    ]))
    story.append(title_tbl)
    story.append(Spacer(1, 0.25*inch))

    # Meta info box
    run_date = datetime.now().strftime("%B %d, %Y  •  %I:%M %p")
    meta_rows = [
        [Paragraph("Prepared By", sty_th),
         Paragraph(prepared_by or "—", sty_td)],
        [Paragraph("Run Date",    sty_th),
         Paragraph(run_date,      sty_td)],
        [Paragraph("Prior Year",  sty_th),
         Paragraph(py,            sty_td)],
        [Paragraph("Current Year",sty_th),
         Paragraph(cy,            sty_td)],
    ]
    meta_tbl = Table(meta_rows, colWidths=[1.6*inch, 5.4*inch])
    meta_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (0,-1), NAVY),
        ("BACKGROUND",    (1,0), (1,-1), LIGHT_BG),
        ("ROWBACKGROUNDS",(1,0), (1,-1), [LIGHT_BG, rl_colors.white]),
        ("GRID",          (0,0), (-1,-1), 0.4, MID_GREY),
        ("TOPPADDING",    (0,0), (-1,-1), 6),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ("LEFTPADDING",   (0,0), (-1,-1), 8),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 0.3*inch))

    # Disclaimer
    disc = ("This analysis was generated using AI (OpenAI o3) and is intended "
            "as a professional review aid. All figures should be verified against "
            "source documents before client delivery or filing.")
    story.append(Paragraph(disc, sty_footer))
    story.append(PageBreak())

    # ══ SECTION HELPER ════════════════════════════════════════════════════
    def section_banner(title_text: str) -> Table:
        t = Table([[Paragraph(title_text, sty_sec_hdr)]], colWidths=[7*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), NAVY),
            ("TOPPADDING", (0,0), (-1,-1), 7),
            ("BOTTOMPADDING",(0,0),(-1,-1), 7),
            ("LEFTPADDING", (0,0), (-1,-1), 10),
        ]))
        return t

    # ══ BODY SECTIONS ═════════════════════════════════════════════════════
    for sec in sections:
        if not sec["body"]:
            continue

        title = sec["title"]
        body  = sec["body"]

        story.append(KeepTogether([
            section_banner(title),
            Spacer(1, 0.06*inch),
        ]))

        # ── Detect pipe table ─────────────────────────────────────────
        table_rows = parse_md_table(body)
        if len(table_rows) >= 2:
            # Separate pre-table narrative
            pre_lines, in_table = [], False
            for line in body.splitlines():
                if "|" in line:
                    in_table = True
                if not in_table:
                    pre_lines.append(line)
            pre_text = "\n".join(pre_lines).strip()
            if pre_text:
                story.append(Paragraph(strip_md(pre_text), sty_body))
                story.append(Spacer(1, 0.05*inch))

            headers = table_rows[0]
            data_rows = table_rows[1:]
            ncols = len(headers)

            # Column widths — last col narrower if "Flag" or "Comment"
            avail = 7.0
            if ncols == 6:  # IS/BS variance tables
                col_w = [2.4, 1.0, 1.0, 1.0, 0.9, 0.7]
            elif ncols == 5:  # Ratio table
                col_w = [1.8, 1.0, 1.0, 1.0, 2.2]
            else:
                w = avail / ncols
                col_w = [w] * ncols

            col_w_in = [c*inch for c in col_w]

            # Build cell data
            pdf_rows = [[Paragraph(h, sty_th) for h in headers]]
            for r in data_rows:
                while len(r) < ncols:
                    r.append("")
                flag = r[-1] if r else ""
                bg   = flag_color_rl(flag)
                row_cells = []
                for i, cell in enumerate(r[:ncols]):
                    align_sty = sty_td_r if i >= 1 and i < ncols-1 else sty_td
                    row_cells.append(Paragraph(strip_md(cell), align_sty))
                pdf_rows.append(row_cells)

            tbl = Table(pdf_rows, colWidths=col_w_in, repeatRows=1)

            # Base style
            tbl_style = [
                ("BACKGROUND",    (0,0), (-1,0),  TEAL),
                ("TEXTCOLOR",     (0,0), (-1,0),  rl_colors.white),
                ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
                ("FONTSIZE",      (0,0), (-1,0),  8.5),
                ("ROWBACKGROUNDS",(0,1), (-1,-1),  [rl_colors.white, LIGHT_BG]),
                ("GRID",          (0,0), (-1,-1),  0.35, MID_GREY),
                ("TOPPADDING",    (0,0), (-1,-1),  4),
                ("BOTTOMPADDING", (0,0), (-1,-1),  4),
                ("LEFTPADDING",   (0,0), (-1,-1),  5),
                ("RIGHTPADDING",  (0,0), (-1,-1),  5),
                ("VALIGN",        (0,0), (-1,-1),  "MIDDLE"),
            ]

            # Per-row flag colouring
            for ri, r in enumerate(data_rows, start=1):
                flag = r[-1] if r else ""
                bg   = flag_color_rl(flag)
                if bg:
                    tbl_style.append(("BACKGROUND", (0,ri), (-1,ri), bg))

            tbl.setStyle(TableStyle(tbl_style))
            story.append(tbl)

        else:
            # ── Numbered or bullet list ──────────────────────────────
            for line in body.splitlines():
                s = line.strip()
                if not s:
                    continue
                s = strip_md(s)
                if re.match(r"^\d+\.", s):
                    story.append(Paragraph(
                        f"<bullet>&bull;</bullet> {s}", sty_bullet))
                elif s.startswith("- ") or s.startswith("• "):
                    story.append(Paragraph(
                        f"<bullet>&bull;</bullet> {s[2:]}", sty_bullet))
                else:
                    story.append(Paragraph(s, sty_body))

        story.append(Spacer(1, 0.12*inch))

    # ── Footer on every page ──────────────────────────────────────────
    def add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(MID_GREY)
        footer_txt = (f"{firm_name}  |  YoY Financial Analysis  |  "
                      f"{py} vs {cy}  |  Page {doc.page}")
        canvas.drawCentredString(letter[0]/2, 0.45*inch, footer_txt)
        canvas.setStrokeColor(TEAL)
        canvas.setLineWidth(0.5)
        canvas.line(0.75*inch, 0.55*inch, letter[0]-0.75*inch, 0.55*inch)
        canvas.restoreState()

    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# WORD GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_row_bg(row, hex_color: str):
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)


def _add_run(para, text: str, bold=False, color_hex: str = None,
             size_pt: int = None):
    run = para.add_run(text)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2],16),
            int(color_hex[2:4],16),
            int(color_hex[4:6],16),
        )
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def build_word(sections: list[dict], cy: str, py: str,
               firm_name: str, prepared_by: str) -> BytesIO:
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────
    for sect in doc.sections:
        sect.left_margin   = Cm(2.0)
        sect.right_margin  = Cm(2.0)
        sect.top_margin    = Cm(2.0)
        sect.bottom_margin = Cm(2.0)

    # ── Styles ───────────────────────────────────────────────────────────
    nstyle = doc.styles["Normal"]
    nstyle.font.name = "Calibri"
    nstyle.font.size = Pt(10)

    # ══ COVER PAGE ═══════════════════════════════════════════════════════
    # Firm banner
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_cell = cover_tbl.cell(0, 0)
    _set_cell_bg(cover_cell, W_NAVY)
    p = cover_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, firm_name.upper(), bold=True, color_hex="FFFFFF", size_pt=13)
    cover_cell.add_paragraph()  # padding

    doc.add_paragraph()

    # Title
    title_tbl = doc.add_table(rows=3, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (txt, sz) in enumerate([
        ("YEAR-OVER-YEAR", 22),
        ("Financial Statement Comparison", 14),
        (f"{py}  →  {cy}", 14),
    ]):
        c = title_tbl.cell(i, 0)
        _set_cell_bg(c, W_TEAL)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(sz)

    doc.add_paragraph()

    # Meta table
    run_date = datetime.now().strftime("%B %d, %Y  —  %I:%M %p")
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    for label, val in [
        ("Prepared By", prepared_by or "—"),
        ("Run Date",    run_date),
        ("Prior Year",  py),
        ("Current Year", cy),
    ]:
        idx = [("Prepared By", "Run Date", "Prior Year", "Current Year")
               .index(label)]
        r = meta.rows[idx[0]]
        _set_cell_bg(r.cells[0], W_NAVY)
        p0 = r.cells[0].paragraphs[0]
        _add_run(p0, label, bold=True, color_hex="FFFFFF", size_pt=9.5)
        _set_cell_bg(r.cells[1], W_LGREY)
        p1 = r.cells[1].paragraphs[0]
        _add_run(p1, val, size_pt=9.5)

    doc.add_paragraph()
    disc = doc.add_paragraph(
        "This analysis was generated using AI (OpenAI o3) and is a professional "
        "review aid only. Verify all figures against source documents before "
        "client delivery or filing.")
    disc.runs[0].font.size = Pt(8)
    disc.runs[0].font.color.rgb = RGBColor(0x90, 0x90, 0x90)
    disc.runs[0].font.italic = True

    doc.add_page_break()

    # ══ BODY SECTIONS ════════════════════════════════════════════════════
    for sec in sections:
        if not sec["body"]:
            continue

        # Section header banner
        hdr_tbl = doc.add_table(rows=1, cols=1)
        hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cell = hdr_tbl.cell(0, 0)
        _set_cell_bg(hdr_cell, W_NAVY)
        hp = hdr_cell.paragraphs[0]
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after  = Pt(0)
        _add_run(hp, sec["title"], bold=True, color_hex="FFFFFF", size_pt=11)

        doc.add_paragraph()

        body = sec["body"]
        table_rows = parse_md_table(body)

        if len(table_rows) >= 2:
            headers   = table_rows[0]
            data_rows = table_rows[1:]
            ncols     = len(headers)

            # Pre-table narrative
            pre_lines, in_table = [], False
            for line in body.splitlines():
                if "|" in line:
                    in_table = True
                if not in_table:
                    pre_lines.append(line)
            pre_text = "\n".join(pre_lines).strip()
            if pre_text:
                doc.add_paragraph(strip_md(pre_text))

            wtbl = doc.add_table(rows=1+len(data_rows), cols=ncols)
            wtbl.style = "Table Grid"
            wtbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Header row
            hrow = wtbl.rows[0]
            _set_row_bg(hrow, W_TEAL)
            for ci, h in enumerate(headers):
                p = hrow.cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, h, bold=True, color_hex="FFFFFF", size_pt=9)

            # Data rows
            for ri, row_data in enumerate(data_rows, start=1):
                while len(row_data) < ncols:
                    row_data.append("")
                flag   = row_data[-1]
                bg_hex = flag_color_word(flag)
                wrow   = wtbl.rows[ri]
                if bg_hex:
                    _set_row_bg(wrow, bg_hex)
                elif ri % 2 == 0:
                    _set_row_bg(wrow, W_LGREY)

                for ci, cell_val in enumerate(row_data[:ncols]):
                    p = wrow.cells[ci].paragraphs[0]
                    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                                   if ci >= 1 and ci < ncols-1
                                   else WD_ALIGN_PARAGRAPH.LEFT)
                    _add_run(p, strip_md(cell_val), size_pt=9)

        else:
            # List / paragraph body
            for line in body.splitlines():
                s = line.strip()
                if not s:
                    doc.add_paragraph()
                    continue
                s = strip_md(s)
                if re.match(r"^\d+\.", s):
                    bp = doc.add_paragraph(style="List Number")
                    _add_run(bp, s, size_pt=10)
                elif s.startswith("- ") or s.startswith("• "):
                    bp = doc.add_paragraph(style="List Bullet")
                    _add_run(bp, s[2:], size_pt=10)
                else:
                    np_ = doc.add_paragraph()
                    _add_run(np_, s, size_pt=10)

        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# ── STREAMLIT PAGE UI ────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

# Custom CSS — keeps branding consistent with the main app
st.markdown("""
<style>
/* ── Sidebar background ── */
[data-testid="stSidebar"] { background: #1B2A4A; }

/* ── ALL sidebar labels → white/light so visible on dark navy ── */
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stTextInput label,
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stTextInput > div > label,
[data-testid="stSidebar"] .stSelectbox > div > label {
    color: #E8EDF5 !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
}

/* ── Sidebar headings → gold accent ── */
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] h2 {
    color: #F0A500 !important;
    font-weight: 700 !important;
}

/* ── Sidebar body paragraphs ── */
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] .stMarkdown p {
    color: #C8D4E8 !important;
}

/* ── Sidebar caption / small text ── */
[data-testid="stSidebar"] small,
[data-testid="stSidebar"] .stCaption,
[data-testid="stSidebar"] .stCaption p {
    color: #8FA0BE !important;
}

/* ── Sidebar input boxes → light background, dark text ── */
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] .stTextInput input {
    color: #1B2A4A !important;
    background: #F4F7FB !important;
    border-radius: 6px !important;
}

/* ── Sidebar selectbox ── */
[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] > div {
    background: #F4F7FB !important;
    color: #1B2A4A !important;
    border-radius: 6px !important;
}

/* ── Sidebar divider ── */
[data-testid="stSidebar"] hr { border-color: #2E4470 !important; }

/* ── Hero banner ── */
.yoy-hero {
    background: linear-gradient(135deg, #1B2A4A 0%, #2E86AB 100%);
    padding: 2rem 2.5rem 1.5rem;
    border-radius: 12px;
    margin-bottom: 1.5rem;
    color: white;
}
.yoy-hero h1 { margin: 0; font-size: 2rem; font-weight: 800; color: white; }
.yoy-hero p  { margin: 0.4rem 0 0; font-size: 1rem; opacity: 0.85; color: white; }

/* ── Section cards ── */
.section-card {
    background: #f4f7fb;
    border-left: 5px solid #2E86AB;
    border-radius: 6px;
    padding: 1rem 1.2rem;
    margin-bottom: 1rem;
}
.metric-box {
    background: white;
    border: 1px solid #D0D7E3;
    border-radius: 8px;
    padding: 0.8rem;
    text-align: center;
}
</style>
""", unsafe_allow_html=True)

# Hero header
st.markdown("""
<div class="yoy-hero">
  <h1>📊 Year-Over-Year Financial Comparison</h1>
  <p>Upload two years of bookkeeping Excel files — AI analyses Income Statement &amp;
     Balance Sheet variances and generates a professional formatted report.</p>
</div>
""", unsafe_allow_html=True)

# ── Sidebar — settings ───────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Settings")

    # API key: read from Streamlit Secrets first (same source as main app),
    # then session_state, then ask manually as last resort
    api_key = ""
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    if not api_key:
        try:
            api_key = st.secrets["openai"]["api_key"]
        except Exception:
            pass
    if not api_key:
        api_key = st.session_state.get("api_key", "")
    if not api_key:
        api_key = st.text_input(
            "OpenAI API Key",
            type="password",
            placeholder="sk-…",
            help="Not found in Streamlit Secrets — enter manually.",
            key="yoy_api_key_input"
        )
    else:
        st.success("✅ API key loaded from Streamlit Secrets")

    model = st.selectbox(
        "AI Model",
        options=["o3", "o4-mini", "gpt-4o"],
        index=0,
        help=(
            "**o3** — OpenAI's most powerful reasoning model (recommended). "
            "Slower but deeper analysis.\n\n"
            "**o4-mini** — Faster reasoning model, still excellent.\n\n"
            "**gpt-4o** — Fast, reliable. No extended reasoning."
        ),
        key="yoy_model"
    )

    firm_name = st.text_input(
        "Firm / Practice Name",
        value="JAINIM CONSULTING INC",
        key="yoy_firm"
    )
    prepared_by = st.text_input(
        "Prepared By",
        value="",
        placeholder="CPA name",
        key="yoy_preparer"
    )

    st.divider()
    st.caption("**Model notes**")
    if model == "o3":
        st.info("o3 uses extended reasoning — expect 30–90 seconds for a thorough analysis.")
    elif model == "o4-mini":
        st.info("o4-mini balances speed and depth — expect 20–50 seconds.")
    else:
        st.info("gpt-4o is fast — expect 15–30 seconds.")

# ── File upload ──────────────────────────────────────────────────────────────
st.markdown("### 📂 Upload Bookkeeping Files")
col_py, col_cy = st.columns(2)

with col_py:
    st.markdown("#### Prior Year")
    py_label = st.text_input("Year label", value="2023", key="yoy_py_lbl")
    py_file  = st.file_uploader("Upload Prior Year Excel", type=["xlsx","xls","xlsm"],
                                 key="yoy_py_file")

with col_cy:
    st.markdown("#### Current Year")
    cy_label = st.text_input("Year label", value="2024", key="yoy_cy_lbl")
    cy_file  = st.file_uploader("Upload Current Year Excel", type=["xlsx","xls","xlsm"],
                                 key="yoy_cy_file")

# ── Sheet selection ──────────────────────────────────────────────────────────
if py_file and cy_file:
    st.markdown("### 🗂️ Select Worksheets")
    py_sheets = get_sheets(py_file)
    cy_sheets = get_sheets(cy_file)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        py_is = st.selectbox(f"IS — {py_label}", py_sheets, key="yoy_py_is")
    with c2:
        py_bs = st.selectbox(f"BS — {py_label}", py_sheets, key="yoy_py_bs")
    with c3:
        cy_is = st.selectbox(f"IS — {cy_label}", cy_sheets, key="yoy_cy_is")
    with c4:
        cy_bs = st.selectbox(f"BS — {cy_label}", cy_sheets, key="yoy_cy_bs")

    # Optional preview
    with st.expander("👁️ Preview sheets before running", expanded=False):
        pc1, pc2 = st.columns(2)
        with pc1:
            st.caption(f"Income Statement — {py_label} ({py_is})")
            st.dataframe(read_sheet(py_file, py_is).head(35), use_container_width=True)
            st.caption(f"Balance Sheet — {py_label} ({py_bs})")
            st.dataframe(read_sheet(py_file, py_bs).head(35), use_container_width=True)
        with pc2:
            st.caption(f"Income Statement — {cy_label} ({cy_is})")
            st.dataframe(read_sheet(cy_file, cy_is).head(35), use_container_width=True)
            st.caption(f"Balance Sheet — {cy_label} ({cy_bs})")
            st.dataframe(read_sheet(cy_file, cy_bs).head(35), use_container_width=True)

    st.divider()

    # ── Run button ────────────────────────────────────────────────────────
    st.markdown("""
    <style>
    div[data-testid="stButton"] > button[kind="primary"] {
        background: linear-gradient(135deg, #1B2A4A, #2E86AB);
        color: white; border: none; border-radius: 8px;
        font-size: 1.1rem; font-weight: 700; padding: 0.8rem;
    }
    </style>
    """, unsafe_allow_html=True)

    run_btn = st.button(
        f"🚀  Run AI Analysis  ({model})",
        type="primary",
        use_container_width=True,
        key="yoy_run"
    )

    if run_btn:
        if not api_key:
            st.error("⚠️ Please enter your OpenAI API key in the sidebar.")
        else:
            prog = st.progress(0, text="Reading Excel sheets…")
            try:
                py_is_df = read_sheet(py_file, py_is)
                py_bs_df = read_sheet(py_file, py_bs)
                cy_is_df = read_sheet(cy_file, cy_is)
                cy_bs_df = read_sheet(cy_file, cy_bs)
                prog.progress(20, text="Sheets loaded — building prompt…")

                prompt = build_prompt(
                    cy_label, py_label,
                    df_to_text(cy_is_df), df_to_text(py_is_df),
                    df_to_text(cy_bs_df), df_to_text(py_bs_df),
                )
                prog.progress(30, text=f"Calling {model} — please wait…")

                raw = call_openai(prompt, api_key, model)
                prog.progress(80, text="Parsing response and building reports…")

                sections = parse_sections(raw)

                # Build PDF
                pdf_buf = build_pdf(
                    sections, cy_label, py_label,
                    firm_name, prepared_by
                )
                # Build Word
                word_buf = build_word(
                    sections, cy_label, py_label,
                    firm_name, prepared_by
                )
                prog.progress(100, text="Done ✅")

                st.session_state["yoy_raw"]      = raw
                st.session_state["yoy_sections"]  = sections
                st.session_state["yoy_pdf"]        = pdf_buf.read()
                st.session_state["yoy_word"]       = word_buf.read()
                st.session_state["yoy_cy"]         = cy_label
                st.session_state["yoy_py"]         = py_label
                st.session_state["yoy_firm"]       = firm_name

            except Exception as e:
                prog.empty()
                st.error(f"❌ Error: {e}")

elif py_file or cy_file:
    st.info("👆 Upload **both** files to continue.")

# ══════════════════════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════════════════════
if "yoy_raw" in st.session_state:
    cy_lbl   = st.session_state["yoy_cy"]
    py_lbl   = st.session_state["yoy_py"]
    sections = st.session_state["yoy_sections"]

    st.success(f"✅ Analysis complete — {py_lbl} vs {cy_lbl}")
    st.divider()

    # ── Download buttons ──────────────────────────────────────────────────
    st.markdown("### 📥 Download Reports")
    dl1, dl2, dl3 = st.columns(3)

    with dl1:
        st.download_button(
            label="📄 Download PDF Report",
            data=st.session_state["yoy_pdf"],
            file_name=f"YoY_{py_lbl}_vs_{cy_lbl}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="yoy_dl_pdf"
        )
    with dl2:
        st.download_button(
            label="📝 Download Word Report",
            data=st.session_state["yoy_word"],
            file_name=f"YoY_{py_lbl}_vs_{cy_lbl}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="yoy_dl_word"
        )
    with dl3:
        st.download_button(
            label="📋 Download Raw Text",
            data=st.session_state["yoy_raw"],
            file_name=f"YoY_{py_lbl}_vs_{cy_lbl}_analysis.txt",
            mime="text/plain",
            use_container_width=True,
            key="yoy_dl_txt"
        )

    st.divider()

    # ── Rendered analysis ─────────────────────────────────────────────────
    st.markdown("### 📊 Analysis Preview")

    for sec in sections:
        if not sec["body"]:
            continue

        # Colour-code section icons
        icon = "📋"
        t = sec["title"].upper()
        if "EXECUTIVE" in t:      icon = "📌"
        elif "INCOME" in t:       icon = "💰"
        elif "BALANCE" in t:      icon = "🏦"
        elif "DRIVER" in t:       icon = "🔑"
        elif "RATIO" in t:        icon = "📐"
        elif "FLAG" in t or "RISK" in t: icon = "🚩"
        elif "CONSISTENCY" in t:  icon = "🔍"
        elif "RECOMMENDATION" in t: icon = "💡"

        with st.expander(f"{icon}  {sec['title']}", expanded=True):
            st.markdown(sec["body"])

    st.divider()
    if st.button("🗑️ Clear Results & Start Over", key="yoy_clear"):
        for k in ["yoy_raw","yoy_sections","yoy_pdf","yoy_word","yoy_cy","yoy_py","yoy_firm"]:
            st.session_state.pop(k, None)
        st.rerun()

elif not (py_file and cy_file):
    # Empty state guide
    st.markdown("---")
    st.markdown("### How it works")
    c1, c2, c3, c4 = st.columns(4)
    for col, num, title, desc in [
        (c1, "1", "Upload Files",   "Upload your Prior Year and Current Year Excel bookkeeping files"),
        (c2, "2", "Select Sheets",  "Choose which sheet is the Income Statement and Balance Sheet for each year"),
        (c3, "3", "Run Analysis",   "o3 reasoning model analyses every account line-by-line"),
        (c4, "4", "Download",       "Get a fully formatted PDF and Word report instantly"),
    ]:
        with col:
            st.markdown(f"""
            <div class="section-card">
              <h3 style="color:#2E86AB;margin:0">Step {num}</h3>
              <h4 style="margin:4px 0">{title}</h4>
              <p style="font-size:0.88rem;color:#555;margin:0">{desc}</p>
            </div>
            """, unsafe_allow_html=True)
