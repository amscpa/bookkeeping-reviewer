"""
pages/4_IS_Analysis.py  —  Income Statement Detailed Analysis
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Reads Bank Statement tab from TWO Excel files.
Filters to INCOME STATEMENT accounts only (row 603 = 'IS').
Sends top 20 transactions per account (with description + client comments).
AI produces a thorough, narrative-style explanation of WHY each account
increased or decreased — reading every description and comment.

WORKBOOK LAYOUT (same as 3_Bank_Analysis.py):
  B2        = Company name
  B4        = Year-end text (year extracted automatically)
  Row 6     = Account short names (col J onwards — col I = Bank, excluded)
  Row 11    = Opening balances
  Rows 12–596 = Transactions  (C=Date, D=Description, E=Client Comments)
  Row 600   = Closing balances
  Row 603   = IS / BS flags  ('IS' = Income Statement account)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import re
from io import BytesIO
from datetime import datetime

import streamlit as st
import pandas as pd
from openai import OpenAI

from reportlab.lib import colors as rl_colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    Table, TableStyle, PageBreak, KeepTogether
)

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="IS Detailed Analysis | Bookkeeping Reviewer",
    page_icon="💰",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY     = rl_colors.HexColor("#1B2A4A")
TEAL     = rl_colors.HexColor("#2E86AB")
GOLD     = rl_colors.HexColor("#F0A500")
LIGHT_BG = rl_colors.HexColor("#F4F7FB")
MID_GREY = rl_colors.HexColor("#D0D7E3")
GREEN_BG = rl_colors.HexColor("#E8F8F0")
RED_BG   = rl_colors.HexColor("#FDECEA")
AMBER_BG = rl_colors.HexColor("#FFF8E1")
GREEN_CL = rl_colors.HexColor("#1A7F4B")
RED_CL   = rl_colors.HexColor("#C0392B")

W_NAVY  = "1B2A4A"
W_TEAL  = "2E86AB"
W_LGREY = "F4F7FB"
W_GREEN = "E8F8F0"
W_RED   = "FDECEA"
W_AMBER = "FFF8E1"

# ── Layout constants (0-based) ────────────────────────────────────────────────
ROW_COMPANY    = 1    # Excel row 2  — company name (col B)
ROW_YEAREND    = 3    # Excel row 4  — year-end text (col B)
ROW_HEADER     = 5    # Excel row 6  — account names
ROW_OPENING    = 10   # Excel row 11 — opening balances
ROW_TXN_START  = 11   # Excel row 12 — first transaction
ROW_TXN_END    = 598  # last transaction row (exclusive)
ROW_CLOSE      = 599  # Excel row 600 — closing balances
ROW_TYPE       = 602  # Excel row 603 — IS / BS flags
COL_B          = 1    # col B
COL_DATE       = 2    # col C — date
COL_DESC       = 3    # col D — description
COL_COMMENTS   = 4    # col E — client comments
COL_ACCT_START = 9    # col J — first account (col I = Bank excluded)
TOP_N          = 20   # ← top 20 transactions per IS account


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def get_sheets(f):
    f.seek(0)
    xl = pd.ExcelFile(BytesIO(f.read()))
    f.seek(0)
    return xl.sheet_names


def extract_year(text):
    m = re.search(r'\b(20\d{2}|19\d{2})\b', str(text))
    return m.group(1) if m else str(text)[:20]


def to_num(v):
    try:
        return float(str(v).replace(",", "").replace("$", "").strip())
    except Exception:
        return 0.0


def extract_is_data(f, sheet_name):
    """
    Returns dict:
      { company, year_label,
        accounts: [{name, opening, closing, net_movement, transactions}] }
    Only IS-flagged accounts are returned.
    Each transaction: {date, desc, comments, amount}
    """
    f.seek(0)
    raw = pd.read_excel(BytesIO(f.read()), sheet_name=sheet_name,
                        header=None, dtype=str)
    f.seek(0)

    # Pad rows if needed
    needed = max(ROW_CLOSE, ROW_TYPE) + 1
    while len(raw) < needed:
        raw.loc[len(raw)] = [""] * len(raw.columns)

    # Company + year
    company    = str(raw.iloc[ROW_COMPANY, COL_B]).strip()
    year_label = extract_year(str(raw.iloc[ROW_YEAREND, COL_B]).strip())
    if company in ("", "nan", "None"):
        company = "Company"

    # Account metadata from row 6 col J+
    acct_names = [str(v).strip() for v in raw.iloc[ROW_HEADER,  COL_ACCT_START:]]
    acct_types = [str(v).strip().upper() for v in raw.iloc[ROW_TYPE,    COL_ACCT_START:]]
    open_vals  = [to_num(v) for v in raw.iloc[ROW_OPENING, COL_ACCT_START:]]
    close_vals = [to_num(v) for v in raw.iloc[ROW_CLOSE,   COL_ACCT_START:]]

    txn_df = raw.iloc[ROW_TXN_START:ROW_TXN_END].copy()

    accounts = []
    for i, name in enumerate(acct_names):
        if not name or name in ("", "nan", "None", "0"):
            continue
        col_idx = COL_ACCT_START + i
        if col_idx >= len(raw.columns):
            break

        # Skip if not an IS account
        flag = acct_types[i] if i < len(acct_types) else ""
        is_income = flag in ("IS", "I", "INCOME", "INCOME STATEMENT")
        if not is_income:
            continue

        series = txn_df.iloc[:, col_idx].apply(to_num)
        net    = series[series.abs() > 0.001].sum()

        opening = open_vals[i]  if i < len(open_vals)  else 0.0
        closing = close_vals[i] if i < len(close_vals) else 0.0

        # Top N transactions by absolute value — keep all descriptions
        mask = series.abs() > 0.001
        sub  = txn_df[mask].copy()
        sub["_amt"] = series[mask].values
        top  = sub.nlargest(TOP_N, "_amt", keep="all").head(TOP_N)

        txns = []
        for _, r in top.iterrows():
            try:
                d = pd.to_datetime(str(r.iloc[COL_DATE]), errors="coerce")
                date_s = d.strftime("%Y-%m-%d") if pd.notna(d) else str(r.iloc[COL_DATE])[:10]
            except Exception:
                date_s = str(r.iloc[COL_DATE])[:10]

            desc = str(r.iloc[COL_DESC]).strip()
            cmt  = str(r.iloc[COL_COMMENTS]).strip()
            # Clean up nan
            if desc in ("nan", "None", ""):
                desc = "—"
            if cmt in ("nan", "None", ""):
                cmt = ""

            txns.append({
                "date":     date_s,
                "desc":     desc[:100],
                "comments": cmt[:100],
                "amount":   r["_amt"],
            })

        accounts.append({
            "name":         name[:50],
            "opening":      opening,
            "closing":      closing,
            "net_movement": net,
            "transactions": txns,
        })

    return {"company": company, "year_label": year_label, "accounts": accounts}


# ══════════════════════════════════════════════════════════════════════════════
# COMPARISON — IS accounts only
# ══════════════════════════════════════════════════════════════════════════════

def build_comparison(cy_data, py_data):
    cy_map = {a["name"]: a for a in cy_data["accounts"]}
    py_map = {a["name"]: a for a in py_data["accounts"]}
    names  = list(dict.fromkeys(
        [a["name"] for a in cy_data["accounts"]] +
        [a["name"] for a in py_data["accounts"]]
    ))
    result = []
    for name in names:
        cy = cy_map.get(name)
        py = py_map.get(name)
        cy_net = cy["net_movement"] if cy else 0.0
        py_net = py["net_movement"] if py else 0.0
        chg    = cy_net - py_net
        pct    = (chg / abs(py_net) * 100) if py_net else (100.0 if cy_net else 0.0)
        result.append({
            "name":          name,
            "cy_opening":    cy["opening"]  if cy else 0.0,
            "cy_closing":    cy["closing"]  if cy else 0.0,
            "py_opening":    py["opening"]  if py else 0.0,
            "py_closing":    py["closing"]  if py else 0.0,
            "cy_net":        cy_net,
            "py_net":        py_net,
            "change_dollar": chg,
            "change_pct":    pct,
            "cy_txns":       cy["transactions"] if cy else [],
            "py_txns":       py["transactions"] if py else [],
        })
    result.sort(key=lambda x: abs(x["change_dollar"]), reverse=True)
    return result


# ══════════════════════════════════════════════════════════════════════════════
# PROMPT — detailed IS narrative
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are a senior Canadian CPA with 20+ years of Alberta small-business experience.
You are writing a DETAILED Income Statement variance analysis for a CPA file review.

CRITICAL RULES — you MUST follow these exactly:

1. For EVERY account, you MUST walk through EACH individual transaction listed and
   explain what it is, what it means, and how it contributed to the variance.
   DO NOT summarize groups of transactions — name each one specifically.

2. Always mention: the exact dollar amount, the date or month, the payee/description,
   and the client comment if one is provided.

3. Explain the STORY of each account: what happened month by month, what was a
   one-time item, what was recurring, what was new this year, what was missing vs last year.

4. Use language like a CPA reviewing a file would use:
   "The $3,800 payment on July 22 to [payee] per client comment '[comment]' represents..."
   "Compared to prior year where the largest item was [payee] at $X, this year..."
   "The spike in [month] is explained by [specific transaction]..."
   "There was a one-time [description] of $X with no prior year equivalent..."

5. NEVER write generic sentences like "expenses increased due to higher costs" or
   "revenue was higher due to increased business activity" — these add no value.
   Every sentence must reference a specific transaction, payee, date, or amount.

6. If a client comment is provided in square brackets, quote it and explain its significance.

7. Compare specific transactions between years — "In {prior_year} the largest item was
   [X] at $Y, whereas in {current_year} the largest item was [A] at $B..."

All amounts are CAD. Be thorough, specific, and actionable."""


def fmt(v):
    return f"${v:,.0f}" if v >= 0 else f"(${abs(v):,.0f})"


def build_prompt(comparison, cy, py):
    lines = [
        "INCOME STATEMENT — DETAILED YEAR-OVER-YEAR ANALYSIS",
        f"Current Year: {cy}    Prior Year: {py}",
        "IS accounts only. Sorted by largest absolute dollar change.",
        f"Top {TOP_N} transactions per account — read EVERY one carefully.",
        "",
    ]

    for a in comparison:
        chg_sign = "INCREASE" if a["change_dollar"] > 0 else "DECREASE"
        lines += [
            "=" * 65,
            f"ACCOUNT: {a['name']}",
            f"  {py} Net: {fmt(a['py_net'])}   {cy} Net: {fmt(a['cy_net'])}",
            f"  CHANGE: {fmt(a['change_dollar'])} ({a['change_pct']:+.1f}%) — {chg_sign}",
            "",
        ]

        if a["cy_txns"]:
            lines.append(f"  ── {cy} TRANSACTIONS (largest {len(a['cy_txns'])}) ──")
            for idx, t in enumerate(a["cy_txns"], 1):
                cmt = f'  CLIENT NOTE: "{t["comments"]}"' if t["comments"] else ""
                lines.append(f"  [{idx:02d}] {t['date']}  {fmt(t['amount']):>10}  {t['desc']}{cmt}")
            lines.append("")

        if a["py_txns"]:
            lines.append(f"  ── {py} TRANSACTIONS (largest {len(a['py_txns'])}) ──")
            for idx, t in enumerate(a["py_txns"], 1):
                cmt = f'  CLIENT NOTE: "{t["comments"]}"' if t["comments"] else ""
                lines.append(f"  [{idx:02d}] {t['date']}  {fmt(t['amount']):>10}  {t['desc']}{cmt}")
            lines.append("")

        if not a["cy_txns"] and not a["py_txns"]:
            lines.append(f"  (No transactions found — account may be new or zero-activity)")
            lines.append("")

    data_block = "\n".join(lines)

    # Build a transaction-by-transaction example to show the AI what level of detail we want
    example_acct = next((a for a in comparison if a["cy_txns"]), None)
    example_str = ""
    if example_acct:
        example_str = f"""
EXAMPLE OF THE LEVEL OF DETAIL REQUIRED for account "{example_acct['name']}":

"The {example_acct['name']} account {'increased' if example_acct['change_dollar'] > 0 else 'decreased'} by \
{fmt(abs(example_acct['change_dollar']))} ({abs(example_acct['change_pct']):.1f}%) from {py} to {cy}. \
""" + (
    f"The largest {cy} transaction was transaction [01] on {example_acct['cy_txns'][0]['date']} "
    f"for {fmt(example_acct['cy_txns'][0]['amount'])} described as '{example_acct['cy_txns'][0]['desc']}'"
    + (f" — the client noted: '{example_acct['cy_txns'][0]['comments']}'" if example_acct['cy_txns'][0]['comments'] else "")
    + ". " if example_acct["cy_txns"] else ""
) + (
    f"In {py} the largest item was {fmt(example_acct['py_txns'][0]['amount'])} on "
    f"{example_acct['py_txns'][0]['date']} for '{example_acct['py_txns'][0]['desc']}'..."
    if example_acct["py_txns"] else ""
) + '"'

    return f"""{data_block}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INSTRUCTIONS — READ CAREFULLY BEFORE WRITING
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{example_str}

For Section 2 (Detailed Account Analysis), you MUST:
- Reference EVERY numbered transaction [01], [02], [03]... listed above
- For each transaction: state the date, amount, what the description tells you,
  and what the client comment (if any) means
- Explain whether each transaction is: recurring (appeared in both years),
  one-time (only in {cy}), increased/decreased vs prior year, or new vendor/payee
- Identify the specific month(s) where activity spiked or dropped
- Compare top transactions between {py} and {cy} by name — do not just say
  "transactions were higher", say WHICH ones were higher and by how much
- If a client comment explains a transaction, incorporate it into your explanation
- Minimum 8-12 sentences per account for accounts with significant variance

PRODUCE EXACTLY THESE SECTIONS:

## 1. EXECUTIVE SUMMARY
5-7 sentences. Name the specific accounts and specific transactions that drove the
overall IS change. Reference actual dollar amounts and payee names from the data above.
DO NOT write generic statements — every sentence must name something specific.

## 2. DETAILED ACCOUNT ANALYSIS

For EVERY account, write a narrative analysis in this exact format:

### [Account Name] — {fmt(0) if not comparison else ''} Change: [$ amount] ([%])

Start with the total change sentence, then walk through EVERY transaction listed:
- Name each transaction by its description and date
- State its amount and whether it is higher, lower, or absent vs prior year
- Reference client comments where provided
- Identify patterns: seasonal, one-time, recurring, new vendor
- End with a one-sentence CPA assessment: Normal / Needs explanation / CRA risk

## 3. INCOME STATEMENT VARIANCE TABLE
Markdown table:
| Account | {py} | {cy} | $ Change | % Change | Assessment |
Assessment: ✅ Normal | ⚠ Needs Explanation | 🔴 CRA Risk | 📈 Revenue Growth | 📉 Revenue Decline

## 4. REVENUE ACCOUNTS — DETAILED COMMENTARY
For revenue/income accounts only:
Walk through each revenue transaction pattern. Are clients the same as last year?
New revenue sources? Lost revenue sources? Any gaps in months? Any unusual revenue timing?

## 5. EXPENSE ACCOUNTS — DETAILED COMMENTARY
For all expense accounts:
Which expenses grew proportionally with revenue (acceptable)?
Which grew disproportionately (needs explanation)?
Which dropped — are receipts possibly missing?
Any descriptions that look like personal expenses coded to business?
Name specific transactions as evidence for each point.

## 6. RED FLAGS & CRA AUDIT RISKS
Numbered list. For each:
- Account + specific transaction(s) from the data above
- Exact dollar amount and date
- Why it is a concern for CRA
- Risk: 🔴 High / 🟡 Medium / 🟢 Low
- Recommended action

## 7. CLIENT MEETING QUESTIONS
8-10 specific questions. Each question must:
- Name the specific account
- Reference the actual transaction description and amount
- Ask for a clear explanation
Example format: "In [account], we see a payment of $X on [date] to [payee description].
Can you confirm what this relates to and provide supporting documentation?"
""".strip()


# ══════════════════════════════════════════════════════════════════════════════
# OPENAI
# ══════════════════════════════════════════════════════════════════════════════

def call_openai(prompt, api_key, model):
    client = OpenAI(api_key=api_key)
    if model.startswith("o"):
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user",   "content": prompt},
            ],
            max_completion_tokens=8000,
        )
    else:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user",   "content": prompt},
            ],
            max_tokens=8000,
            temperature=0.2,
        )
    return resp.choices[0].message.content


# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def parse_sections(text):
    sections, buf, title = [], [], "Preamble"
    for line in text.splitlines():
        if line.startswith("## "):
            if buf:
                sections.append({"title": title, "body": "\n".join(buf).strip()})
            title, buf = line[3:].strip(), []
        else:
            buf.append(line)
    if buf:
        sections.append({"title": title, "body": "\n".join(buf).strip()})
    return sections


def parse_md_table(body):
    rows = []
    for line in body.splitlines():
        s = line.strip()
        if s.startswith("|") and not re.match(r"^\|[-| :]+\|$", s):
            rows.append([c.strip() for c in s.strip("|").split("|")])
    return rows


def strip_md(text):
    return re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", str(text))


def flag_rl(flag):
    if "✅" in flag or "📈" in flag: return GREEN_BG
    if "🔴" in flag or "📉" in flag: return RED_BG
    if "⚠" in flag:                  return AMBER_BG
    return None


def flag_word(flag):
    if "✅" in flag or "📈" in flag: return W_GREEN
    if "🔴" in flag or "📉" in flag: return W_RED
    if "⚠" in flag:                  return W_AMBER
    return None


# ══════════════════════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(sections, comparison, cy, py, firm, preparer):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch)

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sty_wh   = S("wh",  fontSize=11, textColor=rl_colors.white,
                 fontName="Helvetica-Bold", alignment=TA_CENTER)
    sty_wh2  = S("wh2", fontSize=20, textColor=rl_colors.white,
                 fontName="Helvetica-Bold", alignment=TA_CENTER)
    sty_wh3  = S("wh3", fontSize=13, textColor=rl_colors.white,
                 fontName="Helvetica", alignment=TA_CENTER)
    sty_sec  = S("sec", fontSize=12, textColor=rl_colors.white,
                 fontName="Helvetica-Bold", leftIndent=6,
                 spaceAfter=6, spaceBefore=14)
    sty_body = S("bod", fontSize=9.5, textColor=rl_colors.black,
                 fontName="Helvetica", leading=15, spaceAfter=6)
    sty_bul  = S("bul", fontSize=9.5, fontName="Helvetica", leading=14,
                 leftIndent=16, spaceAfter=4, bulletIndent=6)
    sty_th   = S("th",  fontSize=9, textColor=rl_colors.white,
                 fontName="Helvetica-Bold", alignment=TA_CENTER)
    sty_td   = S("td",  fontSize=8.5, fontName="Helvetica",
                 alignment=TA_LEFT, leading=12)
    sty_tdr  = S("tdr", fontSize=8.5, fontName="Helvetica",
                 alignment=TA_RIGHT, leading=12)
    sty_ftr  = S("ftr", fontSize=8, textColor=MID_GREY,
                 alignment=TA_CENTER, fontName="Helvetica-Oblique")

    story = []

    def banner(text, sty, bg, tp=10, bp=10):
        t = Table([[Paragraph(text, sty)]], colWidths=[7*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,-1), bg),
            ("TOPPADDING",    (0,0),(-1,-1), tp),
            ("BOTTOMPADDING", (0,0),(-1,-1), bp),
        ]))
        return t

    # Cover
    story.append(banner(firm.upper(),                     sty_wh,  NAVY, 18, 10))
    story.append(banner("INCOME STATEMENT DETAILED ANALYSIS", sty_wh2, TEAL, 20, 20))
    story.append(banner(f"{py}  →  {cy}  |  IS Accounts Only  |  Top {TOP_N} Transactions",
                        sty_wh3, TEAL, 8, 14))
    story.append(Spacer(1, 0.2*inch))

    run_date = datetime.now().strftime("%B %d, %Y  •  %I:%M %p")
    meta = Table([
        [Paragraph("Prepared By",   sty_th), Paragraph(preparer or "—",            sty_td)],
        [Paragraph("Run Date",      sty_th), Paragraph(run_date,                   sty_td)],
        [Paragraph("Prior Year",    sty_th), Paragraph(py,                         sty_td)],
        [Paragraph("Current Year",  sty_th), Paragraph(cy,                         sty_td)],
        [Paragraph("IS Accounts",   sty_th), Paragraph(str(len(comparison)),        sty_td)],
        [Paragraph("Txns / Account",sty_th), Paragraph(f"Top {TOP_N} by $ amount", sty_td)],
    ], colWidths=[1.6*inch, 5.4*inch])
    meta.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(0,-1), NAVY),
        ("ROWBACKGROUNDS",(1,0),(1,-1), [LIGHT_BG, rl_colors.white]),
        ("GRID",          (0,0),(-1,-1), 0.4, MID_GREY),
        ("TOPPADDING",    (0,0),(-1,-1), 5),
        ("BOTTOMPADDING", (0,0),(-1,-1), 5),
        ("LEFTPADDING",   (0,0),(-1,-1), 8),
    ]))
    story.append(meta)
    story.append(Spacer(1, 0.15*inch))
    story.append(Paragraph(
        "AI analysis powered by OpenAI o3 reasoning model. "
        "Verify all figures against source documents before client delivery or filing.",
        sty_ftr))
    story.append(PageBreak())

    # Quick-ref variance table
    story.append(KeepTogether([
        banner("INCOME STATEMENT VARIANCE QUICK REFERENCE", sty_sec, NAVY, 7, 7),
        Spacer(1, 0.05*inch),
    ]))
    hdr = [Paragraph(h, sty_th) for h in
           ["Account", f"{py} Net", f"{cy} Net", "$ Change", "% Change"]]
    var_rows = [hdr]
    var_ts = [
        ("BACKGROUND",    (0,0),(-1,0),  TEAL),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [rl_colors.white, LIGHT_BG]),
        ("GRID",          (0,0),(-1,-1), 0.35, MID_GREY),
        ("TOPPADDING",    (0,0),(-1,-1), 4),
        ("BOTTOMPADDING", (0,0),(-1,-1), 4),
        ("LEFTPADDING",   (0,0),(-1,-1), 5),
        ("RIGHTPADDING",  (0,0),(-1,-1), 5),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
    ]
    for ri, a in enumerate(comparison, start=1):
        var_rows.append([
            Paragraph(a["name"][:40], sty_td),
            Paragraph(fmt(a["py_net"]),        sty_tdr),
            Paragraph(fmt(a["cy_net"]),        sty_tdr),
            Paragraph(fmt(a["change_dollar"]), sty_tdr),
            Paragraph(f"{a['change_pct']:+.1f}%", sty_tdr),
        ])
        if a["change_dollar"] > 500:
            var_ts.append(("BACKGROUND", (0,ri),(-1,ri), GREEN_BG))
        elif a["change_dollar"] < -500:
            var_ts.append(("BACKGROUND", (0,ri),(-1,ri), RED_BG))

    vt = Table(var_rows,
               colWidths=[2.8*inch, 1.1*inch, 1.1*inch, 1.1*inch, 0.9*inch],
               repeatRows=1)
    vt.setStyle(TableStyle(var_ts))
    story.append(vt)
    story.append(PageBreak())

    # AI sections
    for sec in sections:
        if not sec["body"]:
            continue
        story.append(KeepTogether([
            banner(sec["title"], sty_sec, NAVY, 7, 7),
            Spacer(1, 0.05*inch),
        ]))
        trows = parse_md_table(sec["body"])
        if len(trows) >= 2:
            headers   = trows[0]
            data_rows = trows[1:]
            ncols     = len(headers)
            avail     = 7.0
            if ncols == 6:
                cw = [2.2*inch, 0.9*inch, 0.9*inch, 0.9*inch, 0.8*inch,
                      max(0.3, avail-5.7)*inch]
            else:
                cw = [avail / ncols * inch] * ncols

            pdf_rows = [[Paragraph(h, sty_th) for h in headers]]
            ts = [
                ("BACKGROUND",    (0,0),(-1,0),  TEAL),
                ("ROWBACKGROUNDS",(0,1),(-1,-1), [rl_colors.white, LIGHT_BG]),
                ("GRID",          (0,0),(-1,-1), 0.35, MID_GREY),
                ("TOPPADDING",    (0,0),(-1,-1), 4),
                ("BOTTOMPADDING", (0,0),(-1,-1), 4),
                ("LEFTPADDING",   (0,0),(-1,-1), 5),
                ("RIGHTPADDING",  (0,0),(-1,-1), 5),
                ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
            ]
            for ri, r in enumerate(data_rows, start=1):
                while len(r) < ncols:
                    r.append("")
                bg = flag_rl(r[-1])
                if bg:
                    ts.append(("BACKGROUND", (0,ri),(-1,ri), bg))
                pdf_rows.append([
                    Paragraph(strip_md(cell),
                              sty_tdr if 1 <= ci < ncols-1 else sty_td)
                    for ci, cell in enumerate(r[:ncols])
                ])
            t = Table(pdf_rows, colWidths=cw, repeatRows=1)
            t.setStyle(TableStyle(ts))
            story.append(t)
        else:
            for line in sec["body"].splitlines():
                s = line.strip()
                if not s:
                    continue
                s = strip_md(s)
                if re.match(r"^\d+\.", s) or s.startswith(("- ", "• ")):
                    story.append(Paragraph(
                        f"<bullet>&bull;</bullet> {s.lstrip('0123456789.-• ')}",
                        sty_bul))
                else:
                    story.append(Paragraph(s, sty_body))
        story.append(Spacer(1, 0.1*inch))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(MID_GREY)
        canvas.drawCentredString(
            letter[0]/2, 0.45*inch,
            f"{firm}  |  IS Detailed Analysis  |  {py} vs {cy}  |  Page {doc.page}")
        canvas.setStrokeColor(TEAL)
        canvas.setLineWidth(0.5)
        canvas.line(0.75*inch, 0.55*inch, letter[0]-0.75*inch, 0.55*inch)
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# WORD
# ══════════════════════════════════════════════════════════════════════════════

def _set_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _run(para, text, bold=False, color_hex=None, size_pt=None):
    r = para.add_run(str(text))
    r.bold = bold
    if color_hex:
        r.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16))
    if size_pt:
        r.font.size = Pt(size_pt)
    return r


def build_word(sections, comparison, cy, py, firm, preparer):
    doc = Document()
    for sect in doc.sections:
        sect.left_margin = sect.right_margin = Cm(2.0)
        sect.top_margin  = sect.bottom_margin = Cm(2.0)

    # Cover
    ct = doc.add_table(rows=3, cols=1)
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (txt, bg, sz) in enumerate([
        (firm.upper(),                          W_NAVY, 13),
        ("INCOME STATEMENT DETAILED ANALYSIS",  W_TEAL, 20),
        (f"{py}  →  {cy}  |  IS Accounts Only  |  Top {TOP_N} Transactions", W_TEAL, 12),
    ]):
        c = ct.cell(ri, 0)
        _set_bg(c, bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(sz)

    doc.add_paragraph()

    run_date = datetime.now().strftime("%B %d, %Y  —  %I:%M %p")
    mt = doc.add_table(rows=6, cols=2)
    mt.style = "Table Grid"
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (lbl, val) in enumerate([
        ("Prepared By",   preparer or "—"),
        ("Run Date",      run_date),
        ("Prior Year",    py),
        ("Current Year",  cy),
        ("IS Accounts",   str(len(comparison))),
        ("Txns/Account",  f"Top {TOP_N} by $ amount"),
    ]):
        _set_bg(mt.cell(ri, 0), W_NAVY)
        _run(mt.cell(ri, 0).paragraphs[0], lbl, bold=True, color_hex="FFFFFF", size_pt=9.5)
        _set_bg(mt.cell(ri, 1), W_LGREY)
        _run(mt.cell(ri, 1).paragraphs[0], val, size_pt=9.5)

    doc.add_page_break()

    # Variance quick-ref
    ht = doc.add_table(rows=1, cols=1)
    _set_bg(ht.cell(0, 0), W_NAVY)
    _run(ht.cell(0, 0).paragraphs[0],
         "INCOME STATEMENT VARIANCE QUICK REFERENCE",
         bold=True, color_hex="FFFFFF", size_pt=11)
    doc.add_paragraph()

    vt = doc.add_table(rows=1 + len(comparison), cols=5)
    vt.style = "Table Grid"
    vt.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, h in enumerate(["Account", f"{py} Net", f"{cy} Net", "$ Change", "% Change"]):
        _set_bg(vt.rows[0].cells[ci], W_TEAL)
        p = vt.rows[0].cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, color_hex="FFFFFF", size_pt=9)

    for ri, a in enumerate(comparison, start=1):
        bg  = W_GREEN if a["change_dollar"] > 500 else (
              W_RED   if a["change_dollar"] < -500 else (
              W_LGREY if ri % 2 == 0 else "FFFFFF"))
        row = vt.rows[ri]
        for ci, val in enumerate([
            a["name"][:40],
            fmt(a["py_net"]), fmt(a["cy_net"]),
            fmt(a["change_dollar"]), f"{a['change_pct']:+.1f}%"
        ]):
            _set_bg(row.cells[ci], bg)
            p = row.cells[ci].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if ci >= 1
                           else WD_ALIGN_PARAGRAPH.LEFT)
            _run(p, val, size_pt=9)

    doc.add_page_break()

    # AI sections
    for sec in sections:
        if not sec["body"]:
            continue
        ht2 = doc.add_table(rows=1, cols=1)
        _set_bg(ht2.cell(0, 0), W_NAVY)
        _run(ht2.cell(0, 0).paragraphs[0], sec["title"],
             bold=True, color_hex="FFFFFF", size_pt=11)
        doc.add_paragraph()

        trows = parse_md_table(sec["body"])
        if len(trows) >= 2:
            headers   = trows[0]
            data_rows = trows[1:]
            ncols     = len(headers)
            wt = doc.add_table(rows=1 + len(data_rows), cols=ncols)
            wt.style = "Table Grid"
            wt.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ci, h in enumerate(headers):
                _set_bg(wt.rows[0].cells[ci], W_TEAL)
                p = wt.rows[0].cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, h, bold=True, color_hex="FFFFFF", size_pt=9)
            for ri, rdata in enumerate(data_rows, start=1):
                while len(rdata) < ncols:
                    rdata.append("")
                bg = flag_word(rdata[-1]) or (W_LGREY if ri % 2 == 0 else "FFFFFF")
                for ci, val in enumerate(rdata[:ncols]):
                    _set_bg(wt.rows[ri].cells[ci], bg)
                    p = wt.rows[ri].cells[ci].paragraphs[0]
                    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if 1 <= ci < ncols-1
                                   else WD_ALIGN_PARAGRAPH.LEFT)
                    _run(p, strip_md(val), size_pt=9)
        else:
            for line in sec["body"].splitlines():
                s = line.strip()
                if not s:
                    doc.add_paragraph()
                    continue
                s = strip_md(s)
                if re.match(r"^\d+\.", s):
                    _run(doc.add_paragraph(style="List Number"), s, size_pt=10)
                elif s.startswith(("- ", "• ")):
                    _run(doc.add_paragraph(style="List Bullet"), s[2:], size_pt=10)
                else:
                    _run(doc.add_paragraph(), s, size_pt=10)
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<style>
[data-testid="stSidebar"] { background: #1B2A4A; }
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stTextInput label,
[data-testid="stSidebar"] .stSelectbox label {
    color: #E8EDF5 !important; font-weight: 500 !important; }
[data-testid="stSidebar"] h3 { color: #F0A500 !important; font-weight: 700 !important; }
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] .stMarkdown p { color: #C8D4E8 !important; }
[data-testid="stSidebar"] input {
    color: #1B2A4A !important; background: #F4F7FB !important; border-radius: 6px !important; }
[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] > div {
    background: #F4F7FB !important; color: #1B2A4A !important; }
.is-hero {
    background: linear-gradient(135deg, #1A7F4B 0%, #2E86AB 100%);
    padding: 2rem 2.5rem 1.5rem; border-radius: 12px;
    margin-bottom: 1.5rem; }
.is-hero h1 { margin:0; font-size:2rem; font-weight:800; color:white; }
.is-hero p  { margin:.4rem 0 0; font-size:1rem; opacity:.85; color:white; }
</style>
""", unsafe_allow_html=True)

st.markdown(f"""
<div class="is-hero">
  <h1>💰 Income Statement Detailed Analysis</h1>
  <p>IS accounts only — top {TOP_N} transactions per account with descriptions and client
     comments — AI writes a thorough narrative explanation of every variance.</p>
</div>
""", unsafe_allow_html=True)

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Settings")

    api_key = ""
    for key_path in [["OPENAI_API_KEY"], ["openai", "api_key"]]:
        if api_key:
            break
        try:
            val = st.secrets
            for k in key_path:
                val = val[k]
            api_key = str(val)
        except Exception:
            pass
    if not api_key:
        api_key = st.session_state.get("api_key", "")
    if not api_key:
        api_key = st.text_input("OpenAI API Key", type="password",
                                 placeholder="sk-…", key="is_api_key")
    else:
        st.success("✅ API key loaded")

    model = st.selectbox("AI Model",
        ["o3", "o4-mini", "gpt-4o"], index=0, key="is_model",
        help="o3 = deepest analysis (30-90s). o4-mini = faster. gpt-4o = quickest.")

    firm_name   = st.text_input("Firm Name",   value="JAINIM CONSULTING INC", key="is_firm")
    prepared_by = st.text_input("Prepared By", placeholder="CPA name",        key="is_prep")

    st.divider()
    st.markdown(f"**Top {TOP_N} transactions** per IS account sent to AI")
    st.markdown("Includes both description and client comments")
    if model == "o3":
        st.info("o3: expect 45–120 seconds (detailed analysis)")
    elif model == "o4-mini":
        st.info("o4-mini: expect 25–60 seconds")
    else:
        st.info("gpt-4o: expect 20–40 seconds")

# ── File uploaders ────────────────────────────────────────────────────────────
st.markdown("### 📂 Upload Two Years of Bookkeeping Files")
col_py, col_cy = st.columns(2)

with col_py:
    st.markdown("#### Prior Year")
    py_label = st.text_input("Year label (blank = auto-detect)",
                              value="", key="is_py_lbl")
    py_file  = st.file_uploader("Prior Year Excel (.xlsm/.xlsx)",
                                 type=["xlsx", "xls", "xlsm"], key="is_py_file")

with col_cy:
    st.markdown("#### Current Year")
    cy_label = st.text_input("Year label (blank = auto-detect)",
                              value="", key="is_cy_lbl")
    cy_file  = st.file_uploader("Current Year Excel (.xlsm/.xlsx)",
                                 type=["xlsx", "xls", "xlsm"], key="is_cy_file")

# ── Sheet selectors ───────────────────────────────────────────────────────────
if py_file and cy_file:
    st.markdown("### 🗂️ Select Bank Statement Sheet")
    py_sheets = get_sheets(py_file)
    cy_sheets = get_sheets(cy_file)
    c1, c2 = st.columns(2)
    with c1:
        py_sheet = st.selectbox("Sheet — Prior Year",   py_sheets, key="is_py_sheet")
    with c2:
        cy_sheet = st.selectbox("Sheet — Current Year", cy_sheets, key="is_cy_sheet")

    st.divider()
    run_btn = st.button(
        f"🚀  Run IS Detailed Analysis  ({model})",
        type="primary", use_container_width=True, key="is_run")

    if run_btn:
        if not api_key:
            st.error("⚠️ Please enter your OpenAI API key in the sidebar.")
        else:
            prog = st.progress(0, text="Reading prior year IS accounts…")
            try:
                py_data = extract_is_data(py_file, py_sheet)
                prog.progress(20, text="Reading current year IS accounts…")

                cy_data = extract_is_data(cy_file, cy_sheet)
                n_is    = len(cy_data["accounts"])

                py_lbl  = py_label.strip() or py_data["year_label"]
                cy_lbl  = cy_label.strip() or cy_data["year_label"]
                company = cy_data["company"]

                prog.progress(35, text=f"Found {n_is} IS accounts — building comparison…")
                comparison = build_comparison(cy_data, py_data)

                prog.progress(50, text=f"Calling {model} for detailed narrative analysis…")
                prompt = build_prompt(comparison, cy_lbl, py_lbl)
                raw    = call_openai(prompt, api_key, model)

                prog.progress(88, text="Building PDF and Word reports…")
                sections   = parse_sections(raw)
                pdf_bytes  = build_pdf(sections, comparison, cy_lbl, py_lbl,
                                       company, prepared_by).read()
                word_bytes = build_word(sections, comparison, cy_lbl, py_lbl,
                                        company, prepared_by).read()
                prog.progress(100, text="Done ✅")

                st.session_state.update({
                    "is_raw":        raw,
                    "is_sections":   sections,
                    "is_comparison": comparison,
                    "is_pdf":        pdf_bytes,
                    "is_word":       word_bytes,
                    "is_cy":         cy_lbl,
                    "is_py":         py_lbl,
                    "is_company":    company,
                })

            except Exception as e:
                prog.empty()
                st.error(f"❌ Error: {e}")
                import traceback
                st.code(traceback.format_exc())

elif py_file or cy_file:
    st.info("👆 Upload **both** files to continue.")

# ── Results ───────────────────────────────────────────────────────────────────
if "is_raw" in st.session_state:
    cy_lbl  = st.session_state["is_cy"]
    py_lbl  = st.session_state["is_py"]
    company = st.session_state.get("is_company", "")
    comp    = st.session_state["is_comparison"]
    secs    = st.session_state["is_sections"]

    st.success(
        f"✅  {company}  |  {len(comp)} IS accounts  |  {py_lbl} vs {cy_lbl}")
    st.divider()

    # Downloads
    st.markdown("### 📥 Download Reports")
    d1, d2, d3 = st.columns(3)
    with d1:
        st.download_button("📄 Download PDF",
            st.session_state["is_pdf"],
            f"IS_Analysis_{py_lbl}_vs_{cy_lbl}.pdf",
            "application/pdf",
            use_container_width=True, key="is_dl_pdf")
    with d2:
        st.download_button("📝 Download Word",
            st.session_state["is_word"],
            f"IS_Analysis_{py_lbl}_vs_{cy_lbl}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="is_dl_word")
    with d3:
        st.download_button("📋 Download Text",
            st.session_state["is_raw"],
            f"IS_Analysis_{py_lbl}_vs_{cy_lbl}.txt",
            "text/plain",
            use_container_width=True, key="is_dl_txt")

    st.divider()

    # Variance table
    st.markdown("### 📊 IS Account Variance Summary")
    st.dataframe([{
        "Account":       a["name"],
        f"{py_lbl} Net": f"${a['py_net']:,.0f}",
        f"{cy_lbl} Net": f"${a['cy_net']:,.0f}",
        "$ Change":      f"${a['change_dollar']:+,.0f}",
        "% Change":      f"{a['change_pct']:+.1f}%",
    } for a in comp], use_container_width=True, hide_index=True)

    st.divider()

    # AI sections
    st.markdown("### 🤖 Detailed AI Analysis")
    icons = {
        "EXECUTIVE":  "📌",
        "DETAILED":   "📋",
        "VARIANCE":   "📊",
        "REVENUE":    "💵",
        "EXPENSE":    "💸",
        "RED":        "🚩",
        "CLIENT":     "💬",
    }
    for sec in secs:
        icon = next((v for k, v in icons.items()
                     if k in sec["title"].upper()), "📄")
        with st.expander(f"{icon}  {sec['title']}", expanded=True):
            st.markdown(sec["body"])

    st.divider()
    if st.button("🗑️ Clear Results", key="is_clear"):
        for k in ["is_raw","is_sections","is_comparison",
                  "is_pdf","is_word","is_cy","is_py","is_company"]:
            st.session_state.pop(k, None)
        st.rerun()
