"""
pages/3_Bank_Analysis.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Bank Statement Year-over-Year Analysis — STANDALONE PAGE
• Reads bank statement tab (600+ rows) from TWO separate Excel files
• Aggregates each account column in Python (no raw rows sent to AI)
• Sends compact per-account summary + top transactions to OpenAI o3
• o3 explains WHY each account increased/decreased using transaction descriptions
• Exports formatted PDF + Word report
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

BANK STATEMENT TAB STRUCTURE (hard-coded to match your workbook):
  Col C  (idx 2)  = Date
  Col D  (idx 3)  = Description line 1
  Col E  (idx 4)  = Description line 2
  Col J+ (idx 9+) = Account columns
  Row 9  (idx 8)  = Account names / headers
  Row 11 (idx 10) = Opening balances
  Row 12–599      = Transactions
  Row 600 (idx 599) = Closing balances
  Row 603 (idx 602) = "I" = Income Statement account, "B" = Balance Sheet account
"""

# ─── Standard library ──────────────────────────────────────────────────────
import re
from io import BytesIO
from datetime import datetime

# ─── Third-party ───────────────────────────────────────────────────────────
import streamlit as st
import pandas as pd
import numpy as np
from openai import OpenAI

# ─── ReportLab (PDF) ───────────────────────────────────────────────────────
from reportlab.lib import colors as rl_colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, PageBreak, KeepTogether
)

# ─── python-docx (Word) ────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Bank Statement Analysis | Bookkeeping Reviewer",
    page_icon="🏦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
# COLOURS
# ══════════════════════════════════════════════════════════════════════════════
NAVY     = rl_colors.HexColor("#1B2A4A")
TEAL     = rl_colors.HexColor("#2E86AB")
GOLD     = rl_colors.HexColor("#F0A500")
LIGHT_BG = rl_colors.HexColor("#F4F7FB")
MID_GREY = rl_colors.HexColor("#D0D7E3")
GREEN_BG = rl_colors.HexColor("#E8F8F0")
RED_BG   = rl_colors.HexColor("#FDECEA")
AMBER_BG = rl_colors.HexColor("#FFF8E1")
GREEN_CL = rl_colors.HexColor("#1A7F4B")
RED_CL   = rl_colors.HexColor("#C0392B")

W_NAVY  = "1B2A4A"
W_TEAL  = "2E86AB"
W_LGREY = "F4F7FB"
W_GREEN = "E8F8F0"
W_RED   = "FDECEA"
W_AMBER = "FFF8E1"

# ══════════════════════════════════════════════════════════════════════════════
# LAYOUT CONSTANTS  — adjust here if your workbook ever changes
# ══════════════════════════════════════════════════════════════════════════════
ROW_COMPANY   = 1    # row index (0-based) — row 2  in Excel → cell B2 = company name
ROW_YEAREND   = 3    # row index (0-based) — row 4  in Excel → cell B4 = year-end text
ROW_HEADER    = 5    # row index (0-based) — row 6  in Excel → short account names
ROW_OPENING   = 10   # row index (0-based) — row 11 in Excel → opening balances
ROW_TXN_START = 11   # row index (0-based) — row 12 in Excel → first transaction
ROW_TXN_END   = 598  # row index (0-based) — row 597 inclusive (last real txn ~row 596)
ROW_CLOSE     = 599  # row index (0-based) — row 600 in Excel → closing balances
ROW_TYPE      = 602  # row index (0-based) — row 603 in Excel → 'IS' or 'BS' flags
COL_COMPANY   = 1    # col B (0-based) — company name
COL_DATE      = 2    # col C (0-based) — transaction date
COL_DESC1     = 3    # col D (0-based) — main description
COL_DESC2     = 4    # col E (0-based) — client comments / secondary description
COL_ACCT_START= 9    # col J (0-based) — first account column (skip Bank/cash col I)
TOP_N_TXN     = 8    # top N transactions per account to include in prompt

# ══════════════════════════════════════════════════════════════════════════════
# EXCEL EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def get_sheets(f) -> list:
    f.seek(0)
    xl = pd.ExcelFile(BytesIO(f.read()))
    f.seek(0)
    return xl.sheet_names


def extract_year_from_text(text: str) -> str:
    """Extract 4-digit year from year-end description e.g. 'December 31, 2025' → '2025'."""
    match = re.search(r'\b(20\d{2}|19\d{2})\b', str(text))
    return match.group(1) if match else str(text)[:30]


def extract_bank_data(uploaded_file, sheet_name: str) -> dict:
    """
    Read the bank statement sheet and return:
      {company, year_label, accounts: [{name, acct_type, opening, closing,
                                        net_movement, transactions[]}]}
    Each transaction: {date, desc, amount}
    """
    uploaded_file.seek(0)
    raw = pd.read_excel(
        BytesIO(uploaded_file.read()),
        sheet_name=sheet_name,
        header=None,
        dtype=str,
    )
    uploaded_file.seek(0)

    # ── Pad dataframe if needed ──────────────────────────────────────────
    max_needed_row = max(ROW_CLOSE, ROW_TYPE) + 1
    while len(raw) < max_needed_row:
        raw.loc[len(raw)] = [""] * len(raw.columns)

    # ── Company name from B2 (row index 1, col index 1) ─────────────────
    company = str(raw.iloc[ROW_COMPANY, COL_COMPANY]).strip()
    if company in ("", "nan", "None"):
        company = "Unknown Company"

    # ── Year label from B4 (row index 3, col index 1) ───────────────────
    year_text = str(raw.iloc[ROW_YEAREND, COL_COMPANY]).strip()
    year_label = extract_year_from_text(year_text)

    # ── Account names from Row 6 (index 5), col I onwards (index 8) ─────
    header_row = raw.iloc[ROW_HEADER, COL_ACCT_START:]
    acct_names = [str(v).strip() for v in header_row]

    # ── IS / BS flags from Row 603 (index 602) ───────────────────────────
    type_row   = raw.iloc[ROW_TYPE, COL_ACCT_START:]
    acct_types = [str(v).strip().upper() for v in type_row]

    # ── Opening / closing balances ───────────────────────────────────────
    open_row  = raw.iloc[ROW_OPENING, COL_ACCT_START:]
    close_row = raw.iloc[ROW_CLOSE,   COL_ACCT_START:]

    def to_num(v):
        try:
            return float(str(v).replace(",", "").replace("$", "").strip())
        except Exception:
            return 0.0

    # ── Transaction rows (rows 12–596, index 11–598) ─────────────────────
    txn_df = raw.iloc[ROW_TXN_START:ROW_TXN_END].copy()

    accounts = []
    for i, name in enumerate(acct_names):
        if not name or name in ("", "nan", "None", "0"):
            continue

        col_idx = COL_ACCT_START + i
        if col_idx >= len(raw.columns):
            break

        # Numeric series for this account column
        acct_series = txn_df.iloc[:, col_idx].apply(to_num)

        # Net movement = sum of all posted entries (non-zero only)
        net = acct_series[acct_series != 0].sum()

        # Opening / closing
        opening = to_num(open_row.iloc[i] if i < len(open_row) else 0)
        closing = to_num(close_row.iloc[i] if i < len(close_row) else 0)

        # Top N transactions by absolute value
        non_zero_mask = acct_series.abs() > 0.001
        sub = txn_df[non_zero_mask].copy()
        sub["_amt"] = acct_series[non_zero_mask].values
        sub["_abs"] = sub["_amt"].abs()
        top = sub.nlargest(TOP_N_TXN, "_abs")

        txns = []
        for _, r in top.iterrows():
            raw_date = str(r.iloc[COL_DATE]).strip()
            try:
                d = pd.to_datetime(raw_date, dayfirst=False, errors="coerce")
                date_str = d.strftime("%Y-%m-%d") if pd.notna(d) else raw_date[:10]
            except Exception:
                date_str = raw_date[:10]

            desc1 = str(r.iloc[COL_DESC1]).strip()
            desc2 = str(r.iloc[COL_DESC2]).strip()
            desc  = desc1
            if desc2 and desc2 not in ("nan", "None", ""):
                desc = f"{desc1} ({desc2})"
            txns.append({
                "date":   date_str,
                "desc":   desc[:80],
                "amount": r["_amt"]
            })

        acct_type_val = acct_types[i] if i < len(acct_types) else "?"
        if acct_type_val in ("IS", "I", "INCOME", "INCOME STATEMENT"):
            acct_type = "IS"
        elif acct_type_val in ("BS", "B", "BALANCE", "BALANCE SHEET"):
            acct_type = "BS"
        else:
            acct_type = acct_type_val if acct_type_val != "?" else "IS"

        accounts.append({
            "name":         name[:50],
            "acct_type":    acct_type,
            "opening":      opening,
            "closing":      closing,
            "net_movement": net,
            "transactions": txns,
        })

    return {"company": company, "year_label": year_label, "accounts": accounts}


# ══════════════════════════════════════════════════════════════════════════════
# COMPARISON BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_comparison(cy_data: dict, py_data: dict) -> list:
    """
    Match accounts by name across two years and build a comparison list.
    Accepts dicts returned by extract_bank_data (with 'accounts' key).
    """
    cy_accounts = cy_data["accounts"]
    py_accounts = py_data["accounts"]
    py_map = {a["name"]: a for a in py_accounts}
    cy_map = {a["name"]: a for a in cy_accounts}

    all_names = list(dict.fromkeys(
        [a["name"] for a in cy_accounts] +
        [a["name"] for a in py_accounts]
    ))

    comparison = []
    for name in all_names:
        cy = cy_map.get(name)
        py = py_map.get(name)
        cy_net = cy["net_movement"] if cy else 0.0
        py_net = py["net_movement"] if py else 0.0
        chg    = cy_net - py_net
        pct    = (chg / abs(py_net) * 100) if py_net else (100.0 if cy_net else 0.0)
        acct_type = (cy or py).get("acct_type", "?")
        comparison.append({
            "name":          name,
            "acct_type":     acct_type,
            "cy_opening":    cy["opening"]  if cy else 0.0,
            "cy_closing":    cy["closing"]  if cy else 0.0,
            "py_opening":    py["opening"]  if py else 0.0,
            "py_closing":    py["closing"]  if py else 0.0,
            "cy_net":        cy_net,
            "py_net":        py_net,
            "change_dollar": chg,
            "change_pct":    pct,
            "cy_txns":       cy["transactions"] if cy else [],
            "py_txns":       py["transactions"] if py else [],
        })

    # Sort by absolute dollar change descending
    comparison.sort(key=lambda x: abs(x["change_dollar"]), reverse=True)
    return comparison


# ══════════════════════════════════════════════════════════════════════════════
# PROMPT BUILDER  — compact, token-efficient
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_BANK = """You are a senior Canadian CPA with 20+ years of Alberta small-business experience.
You are reviewing a detailed bookkeeping bank statement comparing two fiscal years.
Each account shows the net movement (sum of all posted entries), opening/closing balances,
and the largest individual transactions with dates and descriptions.
Your job is to explain WHY each account changed — use the transaction descriptions as evidence.
Be specific, reference actual transactions, dates, and amounts. All amounts are CAD."""


def fmt(v: float) -> str:
    return f"${v:,.0f}" if v >= 0 else f"(${abs(v):,.0f})"


def build_bank_prompt(comparison: list, cy_label: str, py_label: str) -> str:
    lines = [
        f"BANK STATEMENT — YEAR-OVER-YEAR ANALYSIS",
        f"Current Year: {cy_label}    Prior Year: {py_label}",
        f"Accounts sorted by largest absolute dollar change (most significant first).",
        "",
    ]

    for a in comparison:
        lines.append(f"{'='*60}")
        lines.append(f"ACCOUNT: {a['name']}  [{a['acct_type']}]")
        lines.append(
            f"  Net Movement:  PY {fmt(a['py_net'])}  →  CY {fmt(a['cy_net'])}"
            f"  |  Change: {fmt(a['change_dollar'])}  ({a['change_pct']:+.1f}%)"
        )
        lines.append(
            f"  Closing Bal:   PY {fmt(a['py_closing'])}  →  CY {fmt(a['cy_closing'])}"
        )

        if a["cy_txns"]:
            lines.append(f"  Top {cy_label} transactions (largest by $):")
            for t in a["cy_txns"]:
                lines.append(f"    {t['date']}  |  {t['desc']}  |  {fmt(t['amount'])}")

        if a["py_txns"]:
            lines.append(f"  Top {py_label} transactions (largest by $):")
            for t in a["py_txns"]:
                lines.append(f"    {t['date']}  |  {t['desc']}  |  {fmt(t['amount'])}")
        lines.append("")

    prompt_data = "\n".join(lines)

    return f"""
Below is a structured account-by-account summary extracted from two years of detailed
bank statement bookkeeping. Each account shows totals and the largest individual
transactions with dates and descriptions.

{prompt_data}

Produce EXACTLY these six sections:

## 1. EXECUTIVE SUMMARY
4–5 sentences: overall picture of what changed most between {py_label} and {cy_label},
the top 2–3 accounts driving the change, and any red flags.

## 2. ACCOUNT-BY-ACCOUNT ANALYSIS
For EVERY account listed above, write a short paragraph (3–6 sentences) that:
- States the dollar and % change clearly
- Identifies the specific transactions driving the change (reference dates and descriptions)
- Explains whether this is a concern, positive trend, or one-time item
- Notes any missing transactions that would be expected

## 3. INCOME STATEMENT ACCOUNTS SUMMARY
Markdown table:
| Account | {py_label} | {cy_label} | $ Change | % Change | Key Driver |
List only Income Statement accounts. Sort by absolute $ change.

## 4. BALANCE SHEET ACCOUNTS SUMMARY
Markdown table:
| Account | {py_label} Closing | {cy_label} Closing | $ Change | % Change | Key Driver |
List only Balance Sheet accounts.

## 5. RED FLAGS & CRA RISK ITEMS
Numbered list. For each: account name, specific concern, risk level
(🔴 High / 🟡 Medium / 🟢 Low), and recommended action.
Focus on: unusual transaction descriptions, round-number entries,
shareholder-related payments, missing expected transactions.

## 6. CLIENT MEETING TALKING POINTS
Numbered list of 4–6 specific questions or observations to raise with the client,
referencing actual transaction descriptions where possible.
""".strip()


# ══════════════════════════════════════════════════════════════════════════════
# OPENAI CALL
# ══════════════════════════════════════════════════════════════════════════════

def call_openai(prompt: str, api_key: str, model: str) -> str:
    client = OpenAI(api_key=api_key)
    if model.startswith("o"):
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_BANK},
                {"role": "user",   "content": prompt},
            ],
            max_completion_tokens=4000,
        )
    else:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_BANK},
                {"role": "user",   "content": prompt},
            ],
            max_tokens=4000,
            temperature=0.2,
        )
    return resp.choices[0].message.content


# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ══════════════════════════════════════════════════════════════════════════════

def parse_sections(text: str) -> list:
    sections, buf, title = [], [], "Preamble"
    for line in text.splitlines():
        if line.startswith("## "):
            if buf:
                sections.append({"title": title, "body": "\n".join(buf).strip()})
            title, buf = line[3:].strip(), []
        else:
            buf.append(line)
    if buf:
        sections.append({"title": title, "body": "\n".join(buf).strip()})
    return sections


def parse_md_table(body: str) -> list:
    rows = []
    for line in body.splitlines():
        s = line.strip()
        if s.startswith("|") and not re.match(r"^\|[-| :]+\|$", s):
            cells = [c.strip() for c in s.strip("|").split("|")]
            rows.append(cells)
    return rows


def strip_md(text: str) -> str:
    return re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", text)


def flag_bg_rl(flag: str):
    if "▲" in flag or "+" in flag:  return GREEN_BG
    if "▼" in flag or "(" in flag:  return RED_BG
    if "⚠" in flag:                 return AMBER_BG
    return None


def flag_bg_word(flag: str) -> str:
    if "▲" in flag or "+" in flag:  return W_GREEN
    if "▼" in flag or "(" in flag:  return W_RED
    if "⚠" in flag:                 return W_AMBER
    return None


# ══════════════════════════════════════════════════════════════════════════════
# PDF BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(sections, comparison, cy, py, firm, preparer) -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch)

    base = getSampleStyleSheet()

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sty_sec = S("Sec", fontSize=12, textColor=rl_colors.white,
                fontName="Helvetica-Bold", leftIndent=6,
                spaceAfter=6, spaceBefore=14)
    sty_body = S("Body", fontSize=9.5, textColor=rl_colors.black,
                 fontName="Helvetica", leading=14, spaceAfter=5)
    sty_bullet = S("Bul", fontSize=9.5, fontName="Helvetica", leading=13,
                   leftIndent=16, spaceAfter=3, bulletIndent=6)
    sty_th = S("TH", fontSize=9, textColor=rl_colors.white,
               fontName="Helvetica-Bold", alignment=TA_CENTER)
    sty_td = S("TD", fontSize=8.5, fontName="Helvetica",
               alignment=TA_LEFT, leading=12)
    sty_td_r = S("TDR", fontSize=8.5, fontName="Helvetica",
                 alignment=TA_RIGHT, leading=12)
    sty_acct_hdr = S("AH", fontSize=10, textColor=rl_colors.white,
                     fontName="Helvetica-Bold", leftIndent=4,
                     spaceAfter=3, spaceBefore=10)
    sty_footer = S("Ftr", fontSize=8, textColor=MID_GREY,
                   alignment=TA_CENTER, fontName="Helvetica-Oblique")

    story = []

    # ── Cover ────────────────────────────────────────────────────────────
    for rows_data in [
        [[firm.upper()], ["BANK STATEMENT ANALYSIS"], [f"{py}  →  {cy}"]],
    ]:
        for i, row in enumerate(rows_data):
            bg   = NAVY if i == 0 else TEAL
            sz   = 11  if i == 0 else (20 if i == 1 else 13)
            bold = i < 2
            t = Table([[Paragraph(row[0],
                         S(f"C{i}", fontSize=sz, textColor=rl_colors.white,
                           fontName="Helvetica-Bold" if bold else "Helvetica",
                           alignment=TA_CENTER))]],
                      colWidths=[7*inch])
            t.setStyle(TableStyle([
                ("BACKGROUND",   (0,0),(-1,-1), bg),
                ("TOPPADDING",   (0,0),(-1,-1), 18 if i==1 else 10),
                ("BOTTOMPADDING",(0,0),(-1,-1), 18 if i==1 else 10),
            ]))
            story.append(t)

    story.append(Spacer(1, 0.2*inch))

    # Meta box
    run_date = datetime.now().strftime("%B %d, %Y  •  %I:%M %p")
    meta = Table([
        [Paragraph("Prepared By",  sty_th), Paragraph(preparer or "—", sty_td)],
        [Paragraph("Run Date",     sty_th), Paragraph(run_date,        sty_td)],
        [Paragraph("Prior Year",   sty_th), Paragraph(py,              sty_td)],
        [Paragraph("Current Year", sty_th), Paragraph(cy,              sty_td)],
        [Paragraph("Accounts",     sty_th), Paragraph(str(len(comparison)), sty_td)],
    ], colWidths=[1.6*inch, 5.4*inch])
    meta.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(0,-1), NAVY),
        ("BACKGROUND",    (1,0),(1,-1), LIGHT_BG),
        ("ROWBACKGROUNDS",(1,0),(1,-1), [LIGHT_BG, rl_colors.white]),
        ("GRID",          (0,0),(-1,-1), 0.4, MID_GREY),
        ("TOPPADDING",    (0,0),(-1,-1), 5),
        ("BOTTOMPADDING", (0,0),(-1,-1), 5),
        ("LEFTPADDING",   (0,0),(-1,-1), 8),
    ]))
    story.append(meta)
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(
        "AI analysis powered by OpenAI o3. Verify all figures against source documents.",
        sty_footer))
    story.append(PageBreak())

    # ── Variance summary table (quick ref) ───────────────────────────────
    t = Table([[Paragraph("ACCOUNT VARIANCE QUICK REFERENCE", sty_sec)]],
              colWidths=[7*inch])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),NAVY),
                            ("TOPPADDING",(0,0),(-1,-1),7),
                            ("BOTTOMPADDING",(0,0),(-1,-1),7),
                            ("LEFTPADDING",(0,0),(-1,-1),10)]))
    story.append(t)
    story.append(Spacer(1, 0.05*inch))

    hdr = [Paragraph(h, sty_th) for h in
           ["Account", "Type", py+" Net", cy+" Net", "$ Change", "% Change"]]
    rows = [hdr]
    for a in comparison:
        pct_str = f"{a['change_pct']:+.1f}%"
        chg_str = fmt(a['change_dollar'])
        bg = (GREEN_BG if a['change_dollar'] > 0
              else RED_BG if a['change_dollar'] < 0 else None)
        rows.append([
            Paragraph(a["name"][:35], sty_td),
            Paragraph(a["acct_type"][:2].upper(), sty_td),
            Paragraph(fmt(a["py_net"]), sty_td_r),
            Paragraph(fmt(a["cy_net"]), sty_td_r),
            Paragraph(chg_str, sty_td_r),
            Paragraph(pct_str, sty_td_r),
        ])

    var_tbl = Table(rows,
        colWidths=[2.4*inch, 0.5*inch, 1.0*inch, 1.0*inch, 1.0*inch, 1.1*inch],
        repeatRows=1)
    style = [
        ("BACKGROUND",    (0,0),(-1,0), TEAL),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white, LIGHT_BG]),
        ("GRID",          (0,0),(-1,-1), 0.35, MID_GREY),
        ("TOPPADDING",    (0,0),(-1,-1), 4),
        ("BOTTOMPADDING", (0,0),(-1,-1), 4),
        ("LEFTPADDING",   (0,0),(-1,-1), 5),
        ("RIGHTPADDING",  (0,0),(-1,-1), 5),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
    ]
    for ri, a in enumerate(comparison, start=1):
        if a["change_dollar"] > 500:
            style.append(("BACKGROUND", (0,ri),(-1,ri), GREEN_BG))
        elif a["change_dollar"] < -500:
            style.append(("BACKGROUND", (0,ri),(-1,ri), RED_BG))
    var_tbl.setStyle(TableStyle(style))
    story.append(var_tbl)
    story.append(PageBreak())

    # ── AI sections ───────────────────────────────────────────────────────
    def sec_banner(title_text):
        t = Table([[Paragraph(title_text, sty_sec)]], colWidths=[7*inch])
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),NAVY),
                                ("TOPPADDING",(0,0),(-1,-1),7),
                                ("BOTTOMPADDING",(0,0),(-1,-1),7),
                                ("LEFTPADDING",(0,0),(-1,-1),10)]))
        return t

    for sec in sections:
        if not sec["body"]:
            continue
        story.append(KeepTogether([sec_banner(sec["title"]),
                                   Spacer(1, 0.05*inch)]))
        body_txt = sec["body"]
        tbl_rows = parse_md_table(body_txt)

        if len(tbl_rows) >= 2:
            headers   = tbl_rows[0]
            data_rows = tbl_rows[1:]
            ncols     = len(headers)
            avail     = 7.0
            cw        = [avail/ncols*inch] * ncols
            if ncols >= 5:
                cw = [2.2*inch, 0.9*inch, 0.9*inch, 0.9*inch, 0.8*inch,
                      max(0.3, avail-5.7)*inch][:ncols]

            pdf_rows = [[Paragraph(h, sty_th) for h in headers]]
            for r in data_rows:
                while len(r) < ncols: r.append("")
                flag = r[-1]
                bg   = flag_bg_rl(flag)
                row_cells = []
                for ci, cell in enumerate(r[:ncols]):
                    st_ = sty_td_r if 1 <= ci < ncols-1 else sty_td
                    row_cells.append(Paragraph(strip_md(cell), st_))
                pdf_rows.append(row_cells)

            t = Table(pdf_rows, colWidths=cw, repeatRows=1)
            ts = [
                ("BACKGROUND",    (0,0),(-1,0),  TEAL),
                ("ROWBACKGROUNDS",(0,1),(-1,-1), [rl_colors.white, LIGHT_BG]),
                ("GRID",          (0,0),(-1,-1), 0.35, MID_GREY),
                ("TOPPADDING",    (0,0),(-1,-1), 4),
                ("BOTTOMPADDING", (0,0),(-1,-1), 4),
                ("LEFTPADDING",   (0,0),(-1,-1), 5),
                ("RIGHTPADDING",  (0,0),(-1,-1), 5),
                ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
            ]
            for ri, r in enumerate(data_rows, start=1):
                bg = flag_bg_rl(r[-1] if r else "")
                if bg:
                    ts.append(("BACKGROUND",(0,ri),(-1,ri),bg))
            t.setStyle(TableStyle(ts))
            story.append(t)
        else:
            for line in body_txt.splitlines():
                s = line.strip()
                if not s: continue
                s = strip_md(s)
                if re.match(r"^\d+\.", s) or s.startswith("- ") or s.startswith("• "):
                    story.append(Paragraph(
                        f"<bullet>&bull;</bullet> {s.lstrip('0123456789.-• ')}",
                        sty_bullet))
                else:
                    story.append(Paragraph(s, sty_body))

        story.append(Spacer(1, 0.1*inch))

    # Footer
    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(MID_GREY)
        canvas.drawCentredString(
            letter[0]/2, 0.45*inch,
            f"{firm}  |  Bank Statement Analysis  |  {py} vs {cy}  |  Page {doc.page}")
        canvas.setStrokeColor(TEAL)
        canvas.setLineWidth(0.5)
        canvas.line(0.75*inch, 0.55*inch, letter[0]-0.75*inch, 0.55*inch)
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# WORD BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _run(para, text, bold=False, color_hex=None, size_pt=None):
    r = para.add_run(str(text))
    r.bold = bold
    if color_hex:
        r.font.color.rgb = RGBColor(
            int(color_hex[0:2],16),
            int(color_hex[2:4],16),
            int(color_hex[4:6],16))
    if size_pt:
        r.font.size = Pt(size_pt)
    return r


def build_word(sections, comparison, cy, py, firm, preparer) -> BytesIO:
    doc = Document()
    for sect in doc.sections:
        sect.left_margin = sect.right_margin = Cm(2.0)
        sect.top_margin  = sect.bottom_margin = Cm(2.0)

    # ── Cover ────────────────────────────────────────────────────────────
    ct = doc.add_table(rows=3, cols=1)
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (txt, bg, sz) in enumerate([
        (firm.upper(),            W_NAVY, 13),
        ("BANK STATEMENT ANALYSIS", W_TEAL, 22),
        (f"{py}  →  {cy}",        W_TEAL, 14),
    ]):
        c = ct.cell(ri, 0)
        _set_bg(c, bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r.font.size = Pt(sz)

    doc.add_paragraph()

    run_date = datetime.now().strftime("%B %d, %Y  —  %I:%M %p")
    mt = doc.add_table(rows=5, cols=2)
    mt.style = "Table Grid"
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (lbl, val) in enumerate([
        ("Prepared By",  preparer or "—"),
        ("Run Date",     run_date),
        ("Prior Year",   py),
        ("Current Year", cy),
        ("Accounts",     str(len(comparison))),
    ]):
        _set_bg(mt.cell(ri,0), W_NAVY)
        _run(mt.cell(ri,0).paragraphs[0], lbl, bold=True, color_hex="FFFFFF", size_pt=9.5)
        _set_bg(mt.cell(ri,1), W_LGREY)
        _run(mt.cell(ri,1).paragraphs[0], val, size_pt=9.5)

    doc.add_page_break()

    # ── Variance quick-ref table ──────────────────────────────────────────
    ht = doc.add_table(rows=1, cols=1)
    ht.alignment = WD_TABLE_ALIGNMENT.LEFT
    hc = ht.cell(0,0)
    _set_bg(hc, W_NAVY)
    _run(hc.paragraphs[0], "ACCOUNT VARIANCE QUICK REFERENCE",
         bold=True, color_hex="FFFFFF", size_pt=11)

    doc.add_paragraph()

    vt = doc.add_table(rows=1+len(comparison), cols=6)
    vt.style = "Table Grid"
    vt.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdrs = ["Account","Type", f"{py} Net", f"{cy} Net", "$ Change", "% Change"]
    hrow = vt.rows[0]
    for ci, h in enumerate(hdrs):
        _set_bg(hrow.cells[ci], W_TEAL)
        p = hrow.cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, color_hex="FFFFFF", size_pt=9)

    for ri, a in enumerate(comparison, start=1):
        row = vt.rows[ri]
        bg  = W_GREEN if a["change_dollar"] > 500 else (W_RED if a["change_dollar"] < -500 else (W_LGREY if ri%2==0 else "FFFFFF"))
        for ci in range(6):
            _set_bg(row.cells[ci], bg)
        vals = [a["name"][:35], a["acct_type"][:2].upper(),
                fmt(a["py_net"]), fmt(a["cy_net"]),
                fmt(a["change_dollar"]), f"{a['change_pct']:+.1f}%"]
        for ci, v in enumerate(vals):
            p = row.cells[ci].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if ci >= 2
                           else WD_ALIGN_PARAGRAPH.LEFT)
            _run(p, v, size_pt=9)

    doc.add_page_break()

    # ── AI sections ───────────────────────────────────────────────────────
    for sec in sections:
        if not sec["body"]: continue

        ht2 = doc.add_table(rows=1, cols=1)
        hc2 = ht2.cell(0,0)
        _set_bg(hc2, W_NAVY)
        _run(hc2.paragraphs[0], sec["title"], bold=True, color_hex="FFFFFF", size_pt=11)
        doc.add_paragraph()

        body  = sec["body"]
        trows = parse_md_table(body)
        if len(trows) >= 2:
            hdrs2     = trows[0]
            data_rows = trows[1:]
            ncols     = len(hdrs2)
            wt = doc.add_table(rows=1+len(data_rows), cols=ncols)
            wt.style = "Table Grid"
            wt.alignment = WD_TABLE_ALIGNMENT.LEFT
            hr2 = wt.rows[0]
            for ci, h in enumerate(hdrs2):
                _set_bg(hr2.cells[ci], W_TEAL)
                p = hr2.cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, h, bold=True, color_hex="FFFFFF", size_pt=9)
            for ri, rdata in enumerate(data_rows, start=1):
                while len(rdata) < ncols: rdata.append("")
                flag = rdata[-1]
                bg   = flag_bg_word(flag) or ("F4F7FB" if ri%2==0 else "FFFFFF")
                wr   = wt.rows[ri]
                for ci, val in enumerate(rdata[:ncols]):
                    _set_bg(wr.cells[ci], bg)
                    p = wr.cells[ci].paragraphs[0]
                    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if 1 <= ci < ncols-1
                                   else WD_ALIGN_PARAGRAPH.LEFT)
                    _run(p, strip_md(val), size_pt=9)
        else:
            for line in body.splitlines():
                s = line.strip()
                if not s:
                    doc.add_paragraph()
                    continue
                s = strip_md(s)
                if re.match(r"^\d+\.", s):
                    bp = doc.add_paragraph(style="List Number")
                    _run(bp, s, size_pt=10)
                elif s.startswith("- ") or s.startswith("• "):
                    bp = doc.add_paragraph(style="List Bullet")
                    _run(bp, s[2:], size_pt=10)
                else:
                    np_ = doc.add_paragraph()
                    _run(np_, s, size_pt=10)

        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# ── STREAMLIT UI ─────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<style>
[data-testid="stSidebar"] { background: #1B2A4A; }
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stTextInput label,
[data-testid="stSidebar"] .stSelectbox label {
    color: #E8EDF5 !important; font-weight: 500 !important; }
[data-testid="stSidebar"] h3 { color: #F0A500 !important; font-weight: 700 !important; }
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] .stMarkdown p { color: #C8D4E8 !important; }
[data-testid="stSidebar"] input { color: #1B2A4A !important;
    background: #F4F7FB !important; border-radius: 6px !important; }
[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] > div {
    background: #F4F7FB !important; color: #1B2A4A !important; }
.bank-hero {
    background: linear-gradient(135deg, #1B2A4A 0%, #2E86AB 100%);
    padding: 2rem 2.5rem 1.5rem; border-radius: 12px;
    margin-bottom: 1.5rem; color: white; }
.bank-hero h1 { margin:0; font-size:2rem; font-weight:800; color:white; }
.bank-hero p  { margin:.4rem 0 0; font-size:1rem; opacity:.85; color:white; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="bank-hero">
  <h1>🏦 Bank Statement Year-over-Year Analysis</h1>
  <p>Upload two bookkeeping files — Python aggregates every account column,
     then AI explains why each account increased or decreased using
     actual transaction descriptions.</p>
</div>
""", unsafe_allow_html=True)

# ── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Settings")

    api_key = ""
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    if not api_key:
        try:
            api_key = st.secrets["openai"]["api_key"]
        except Exception:
            pass
    if not api_key:
        api_key = st.session_state.get("api_key", "")
    if not api_key:
        api_key = st.text_input("OpenAI API Key", type="password",
                                 placeholder="sk-…", key="bank_api_key")
    else:
        st.success("✅ API key loaded")

    model = st.selectbox("AI Model",
        options=["o3","o4-mini","gpt-4o"], index=0, key="bank_model",
        help="o3 = best analysis. o4-mini = faster. gpt-4o = quickest.")

    firm_name   = st.text_input("Firm Name", value="JAINIM CONSULTING INC", key="bank_firm")
    prepared_by = st.text_input("Prepared By", placeholder="CPA name", key="bank_prep")

    st.divider()
    st.markdown("**Layout settings** *(adjust if your workbook changes)*")
    acct_start_row = st.number_input("Account header row", value=9, min_value=1,
        help="Row number in Excel that contains account names (default: row 9)")
    acct_start_col = st.number_input("First account column number",
        value=10, min_value=1,
        help="Column J = 10 (first account after Bank col I which is excluded)")
    opening_row = st.number_input("Opening balance row", value=11, min_value=1)
    closing_row = st.number_input("Closing balance row", value=600, min_value=1)
    type_row_num = st.number_input("IS/BS flag row",     value=603, min_value=1)
    top_n = st.number_input("Top transactions per account", value=8,
                             min_value=3, max_value=15)

    if model == "o3":
        st.info("o3: expect 30–90 seconds")
    elif model == "o4-mini":
        st.info("o4-mini: expect 20–50 seconds")
    else:
        st.info("gpt-4o: expect 15–30 seconds")

# ── File uploaders ────────────────────────────────────────────────────────────
st.markdown("### 📂 Upload Bookkeeping Files")
col_py, col_cy = st.columns(2)
with col_py:
    st.markdown("#### Prior Year")
    py_label = st.text_input("Year label", value="2023", key="bank_py_lbl")
    py_file  = st.file_uploader("Prior Year Excel",
                                 type=["xlsx","xls","xlsm"], key="bank_py_file")
with col_cy:
    st.markdown("#### Current Year")
    cy_label = st.text_input("Year label", value="2024", key="bank_cy_lbl")
    cy_file  = st.file_uploader("Current Year Excel",
                                 type=["xlsx","xls","xlsm"], key="bank_cy_file")

if py_file and cy_file:
    st.markdown("### 🗂️ Select Bank Statement Sheet")
    py_sheets = get_sheets(py_file)
    cy_sheets = get_sheets(cy_file)
    c1, c2 = st.columns(2)
    with c1:
        py_sheet = st.selectbox(f"Sheet — {py_label}", py_sheets, key="bank_py_sheet")
    with c2:
        cy_sheet = st.selectbox(f"Sheet — {cy_label}", cy_sheets, key="bank_cy_sheet")

    st.divider()

    run_btn = st.button(f"🚀  Run Bank Statement Analysis  ({model})",
                        type="primary", use_container_width=True, key="bank_run")

    if run_btn:
        if not api_key:
            st.error("⚠️ Please enter your OpenAI API key in the sidebar.")
        else:
            prog = st.progress(0, text="Reading prior year bank statement…")
            try:
                # Override layout constants from sidebar inputs
                global ROW_HEADER, COL_ACCT_START, ROW_OPENING, ROW_CLOSE
                global ROW_TYPE, TOP_N_TXN
                ROW_HEADER     = int(acct_start_row)  - 1
                COL_ACCT_START = int(acct_start_col)  - 1
                ROW_OPENING    = int(opening_row)      - 1
                ROW_CLOSE      = int(closing_row)      - 1
                ROW_TYPE       = int(type_row_num)     - 1
                TOP_N_TXN      = int(top_n)

                py_data = extract_bank_data(py_file, py_sheet)
                prog.progress(25, text="Reading current year bank statement…")

                cy_data = extract_bank_data(cy_file, cy_sheet)
                n_accts = len(cy_data["accounts"])

                # Auto-populate year labels from B4 if user left defaults
                cy_lbl_auto = cy_data["year_label"] if cy_label in ("2024","") else cy_label
                py_lbl_auto = py_data["year_label"] if py_label in ("2023","") else py_label
                company     = cy_data["company"]

                prog.progress(40, text=f"Extracted {n_accts} accounts — building comparison…")

                comparison = build_comparison(cy_data, py_data)
                prog.progress(55, text=f"Calling {model} — this may take up to 90 seconds…")

                prompt = build_bank_prompt(comparison, cy_lbl_auto, py_lbl_auto)
                raw    = call_openai(prompt, api_key, model)
                prog.progress(90, text="Building PDF and Word reports…")

                sections = parse_sections(raw)
                pdf_buf  = build_pdf(sections, comparison, cy_lbl_auto, py_lbl_auto,
                                     company, prepared_by)
                word_buf = build_word(sections, comparison, cy_lbl_auto, py_lbl_auto,
                                      company, prepared_by)
                prog.progress(100, text="Done ✅")

                st.session_state["bank_raw"]        = raw
                st.session_state["bank_sections"]   = sections
                st.session_state["bank_comparison"] = comparison
                st.session_state["bank_pdf"]        = pdf_buf.read()
                st.session_state["bank_word"]       = word_buf.read()
                st.session_state["bank_cy"]         = cy_lbl_auto
                st.session_state["bank_py"]         = py_lbl_auto
                st.session_state["bank_company"]    = company
                prog.progress(100, text="Done ✅")

                st.session_state["bank_raw"]        = raw
                st.session_state["bank_sections"]   = sections
                st.session_state["bank_comparison"] = comparison
                st.session_state["bank_pdf"]        = pdf_buf.read()
                st.session_state["bank_word"]       = word_buf.read()
                st.session_state["bank_cy"]         = cy_label
                st.session_state["bank_py"]         = py_label

            except Exception as e:
                prog.empty()
                st.error(f"❌ Error: {e}")
                import traceback
                st.code(traceback.format_exc())

elif py_file or cy_file:
    st.info("👆 Please upload **both** files to continue.")

# ── Results ───────────────────────────────────────────────────────────────────
if "bank_raw" in st.session_state:
    cy_lbl  = st.session_state["bank_cy"]
    py_lbl  = st.session_state["bank_py"]
    company = st.session_state.get("bank_company", "")
    comp    = st.session_state["bank_comparison"]
    secs    = st.session_state["bank_sections"]

    st.success(f"✅ {company}  |  {len(comp)} accounts compared  |  {py_lbl} vs {cy_lbl}")
    st.divider()

    # Downloads
    st.markdown("### 📥 Download Reports")
    d1, d2, d3 = st.columns(3)
    with d1:
        st.download_button("📄 Download PDF", st.session_state["bank_pdf"],
            f"BankAnalysis_{py_lbl}_vs_{cy_lbl}.pdf", "application/pdf",
            use_container_width=True, key="bank_dl_pdf")
    with d2:
        st.download_button("📝 Download Word", st.session_state["bank_word"],
            f"BankAnalysis_{py_lbl}_vs_{cy_lbl}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="bank_dl_word")
    with d3:
        st.download_button("📋 Download Text", st.session_state["bank_raw"],
            f"BankAnalysis_{py_lbl}_vs_{cy_lbl}.txt", "text/plain",
            use_container_width=True, key="bank_dl_txt")

    st.divider()

    # Quick variance table
    st.markdown("### 📊 Account Variance Summary")
    tbl_data = []
    for a in comp:
        tbl_data.append({
            "Account":   a["name"],
            "Type":      a["acct_type"],
            f"{py_lbl} Net": f"${a['py_net']:,.0f}",
            f"{cy_lbl} Net": f"${a['cy_net']:,.0f}",
            "$ Change":  f"${a['change_dollar']:+,.0f}",
            "% Change":  f"{a['change_pct']:+.1f}%",
        })
    st.dataframe(tbl_data, use_container_width=True, hide_index=True)

    st.divider()

    # AI analysis sections
    st.markdown("### 🤖 AI Analysis")
    icons = {"EXECUTIVE":"📌","ACCOUNT":"📋","INCOME":"💰",
             "BALANCE":"🏦","RED":"🚩","CLIENT":"💡"}
    for sec in secs:
        icon = next((v for k,v in icons.items() if k in sec["title"].upper()), "📄")
        with st.expander(f"{icon}  {sec['title']}", expanded=True):
            st.markdown(sec["body"])

    st.divider()
    if st.button("🗑️ Clear Results", key="bank_clear"):
        for k in ["bank_raw","bank_sections","bank_comparison",
                  "bank_pdf","bank_word","bank_cy","bank_py"]:
            st.session_state.pop(k, None)
        st.rerun()
